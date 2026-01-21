import base64
import io
import json
import os
import re
import uuid
import smtplib
from email.message import EmailMessage
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone, date, time
from typing import List, Dict, Any, Optional, Tuple

import fitz  # PyMuPDF
from PIL import Image
import streamlit as st

# --- Optional OpenAI (kept for PDF parsing flow) ---
try:
    from openai import OpenAI
except Exception:
    OpenAI = None  # type: ignore

from graph_client import GraphClient, GraphConfig, GraphAPIError, GraphAuthError
from audit_log import AuditLog, LogLevel, log_structured, InterviewStatus
from ics_utils import ICSInvite, stable_uid, ICSValidationError, create_ics_from_interview, generate_cancellation_ics
from timezone_utils import to_utc, from_utc, iso_utc, is_valid_timezone, safe_zoneinfo
from export_utils import (
    export_interviews_csv,
    export_audit_log_csv,
    format_audit_entry_human,
    filter_interviews_for_export,
    filter_audit_entries,
    AUDIT_ACTION_DESCRIPTIONS,
)


# ----------------------------
# Input Validation
# ----------------------------
import re as _re
from typing import Tuple as _Tuple

# Email regex (RFC 5322 simplified)
_EMAIL_REGEX = _re.compile(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$')

# Date/time patterns from OpenAI output
_DATE_REGEX = _re.compile(r'^\d{4}-\d{2}-\d{2}$')
_TIME_REGEX = _re.compile(r'^\d{2}:\d{2}$')


class ValidationError(ValueError):
    """Raised when input validation fails."""
    def __init__(self, field: str, message: str):
        self.field = field
        self.message = message
        super().__init__(f"{field}: {message}")


# Maximum number of candidates allowed in a single batch
MAX_CANDIDATES = 20


@dataclass
class CandidateValidationResult:
    """Result of validating a single candidate entry."""
    original: str           # Original input string
    email: Optional[str]    # Validated email (None if invalid)
    name: str               # Parsed name (empty string if not provided)
    is_valid: bool          # True if email is valid
    error: Optional[str]    # Error message if invalid


@dataclass
class SchedulingResult:
    """Result of scheduling an interview for one candidate."""
    candidate_email: str
    candidate_name: str
    success: bool
    event_id: Optional[str]
    teams_url: Optional[str]
    error: Optional[str]


@dataclass
class CompanyConfig:
    """Company branding configuration for emails."""
    name: str
    logo_url: Optional[str]
    primary_color: str
    website: Optional[str]
    sender_email: str

    @property
    def signature_name(self) -> str:
        """Return the company signature name for emails."""
        return f"{self.name} Talent Acquisition Team"


@dataclass
class LayoutConfig:
    """Layout and display configuration for UI branding."""
    show_sidebar: bool
    show_footer: bool
    show_powered_by: bool
    header_style: str  # "full", "compact", "minimal"


def validate_email(email: str, field_name: str = "email") -> str:
    """Validate email format. Returns cleaned email or raises ValidationError."""
    if not email:
        raise ValidationError(field_name, "Email is required")
    email = email.strip().lower()
    if not _EMAIL_REGEX.match(email):
        raise ValidationError(field_name, f"Invalid email format: {email}")
    if len(email) > 254:  # RFC 5321 limit
        raise ValidationError(field_name, "Email too long (max 254 characters)")
    return email


def validate_email_optional(email: Optional[str], field_name: str = "email") -> Optional[str]:
    """Validate email if provided, return None if empty."""
    if not email or not email.strip():
        return None
    return validate_email(email, field_name)


# Pattern for "Name <email>" format
_NAME_EMAIL_PATTERN = _re.compile(r'^(.+?)\s*<([^>]+)>$')


def _parse_single_candidate(entry: str) -> CandidateValidationResult:
    """
    Parse a single candidate entry.
    Supports formats: 'email@example.com' or 'Name <email@example.com>'
    """
    entry = entry.strip()
    if not entry:
        return CandidateValidationResult(
            original=entry,
            email=None,
            name="",
            is_valid=False,
            error="Empty entry"
        )

    # Try to match "Name <email>" format
    match = _NAME_EMAIL_PATTERN.match(entry)
    if match:
        name = match.group(1).strip()
        email_raw = match.group(2).strip()
    else:
        name = ""
        email_raw = entry

    try:
        validated_email = validate_email(email_raw, "candidate email")
        return CandidateValidationResult(
            original=entry,
            email=validated_email,
            name=name,
            is_valid=True,
            error=None
        )
    except ValidationError as e:
        return CandidateValidationResult(
            original=entry,
            email=None,
            name=name,
            is_valid=False,
            error=e.message
        )


def parse_candidate_emails(raw_input: str) -> List[CandidateValidationResult]:
    """
    Parse semicolon-separated candidate emails with optional names.

    Formats supported:
    - "email@example.com"
    - "Name <email@example.com>"
    - "email1@example.com; email2@example.com"

    Returns list of CandidateValidationResult objects with validation status.
    Enforces MAX_CANDIDATES limit and detects duplicates.
    """
    results: List[CandidateValidationResult] = []
    if not raw_input or not raw_input.strip():
        return results

    entries = [e.strip() for e in raw_input.split(';') if e.strip()]

    # Check for exceeding limit
    if len(entries) > MAX_CANDIDATES:
        # Parse first MAX_CANDIDATES normally, mark rest as exceeding limit
        for i, entry in enumerate(entries):
            if i < MAX_CANDIDATES:
                results.append(_parse_single_candidate(entry))
            else:
                results.append(CandidateValidationResult(
                    original=entry,
                    email=None,
                    name="",
                    is_valid=False,
                    error=f"Exceeds maximum of {MAX_CANDIDATES} candidates"
                ))
        return results

    # Track seen emails for duplicate detection
    seen_emails: set = set()

    for entry in entries:
        result = _parse_single_candidate(entry)

        # Check for duplicates among valid entries
        if result.is_valid and result.email:
            if result.email in seen_emails:
                results.append(CandidateValidationResult(
                    original=result.original,
                    email=None,
                    name=result.name,
                    is_valid=False,
                    error="Duplicate email address"
                ))
            else:
                seen_emails.add(result.email)
                results.append(result)
        else:
            results.append(result)

    return results


def validate_slot(slot: dict) -> _Tuple[str, str, str]:
    """Validate slot dict from OpenAI parsing. Returns (date, start, end) tuple."""
    if not isinstance(slot, dict):
        raise ValidationError("slot", "Slot must be a dictionary")

    date = slot.get("date", "")
    start = slot.get("start", "")
    end = slot.get("end", "")

    if not _DATE_REGEX.match(date):
        raise ValidationError("slot.date", f"Invalid date format: {date}. Expected YYYY-MM-DD")
    if not _TIME_REGEX.match(start):
        raise ValidationError("slot.start", f"Invalid start time format: {start}. Expected HH:MM")
    if end and not _TIME_REGEX.match(end):
        raise ValidationError("slot.end", f"Invalid end time format: {end}. Expected HH:MM")

    return date, start, end


# ----------------------------
# Configuration helpers
# ----------------------------
def get_secret(key: str, default: Any = None) -> Any:
    # st.secrets behaves like a dict on Streamlit Cloud; local dev can use env vars too.
    try:
        if key in st.secrets:
            return st.secrets.get(key, default)
    except Exception:
        pass
    return os.getenv(key.upper(), default)


def get_default_timezone() -> str:
    return get_secret("default_timezone", "Europe/London")


def get_audit_log_path() -> str:
    return get_secret("audit_log_path", "audit_log.db")


def _get_branding_settings_path() -> str:
    """Get path for persistent branding settings file."""
    return get_secret("branding_settings_path", "branding_settings.json")


def _load_branding_settings() -> Dict[str, Any]:
    """Load branding settings from persistent storage."""
    path = _get_branding_settings_path()
    try:
        if os.path.exists(path):
            with open(path, 'r') as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def _save_branding_settings(settings: Dict[str, Any]) -> None:
    """Save branding settings to persistent storage."""
    path = _get_branding_settings_path()
    try:
        with open(path, 'w') as f:
            json.dump(settings, f, indent=2)
    except Exception as e:
        st.warning(f"Could not save branding settings: {e}")


def get_graph_config() -> Optional[GraphConfig]:
    tenant_id = get_secret("graph_tenant_id")
    client_id = get_secret("graph_client_id")
    client_secret = get_secret("graph_client_secret")
    scheduler_mailbox = get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com")

    if not (tenant_id and client_id and client_secret and scheduler_mailbox):
        return None
    return GraphConfig(
        tenant_id=str(tenant_id),
        client_id=str(client_id),
        client_secret=str(client_secret),
        scheduler_mailbox=str(scheduler_mailbox),
    )


def get_company_config() -> CompanyConfig:
    """
    Load company branding configuration.
    Checks session state for user overrides first, then falls back to secrets.
    """
    # Check session state for user customizations (if initialized)
    custom_name = st.session_state.get("custom_company_name") if hasattr(st, "session_state") else None
    custom_logo = st.session_state.get("custom_logo_data") if hasattr(st, "session_state") else None
    custom_color = st.session_state.get("custom_primary_color") if hasattr(st, "session_state") else None

    return CompanyConfig(
        name=custom_name or get_secret("company_name", "PowerDash HR"),
        logo_url=custom_logo or get_secret("company_logo_url"),
        primary_color=custom_color or get_secret("company_primary_color", "#0066CC"),
        website=get_secret("company_website"),
        sender_email=get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"),
    )


def get_layout_config() -> LayoutConfig:
    """Load layout configuration from secrets for UI branding."""
    return LayoutConfig(
        show_sidebar=get_secret("show_sidebar", False),
        show_footer=get_secret("show_footer", True),
        show_powered_by=get_secret("show_powered_by", True),
        header_style=get_secret("header_style", "full"),
    )


def _lighten_color(hex_color: str, factor: float) -> str:
    """Lighten a hex color by mixing with white."""
    hex_color = hex_color.lstrip('#')
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    r = int(r + (255 - r) * factor)
    g = int(g + (255 - g) * factor)
    b = int(b + (255 - b) * factor)
    return f"#{r:02x}{g:02x}{b:02x}"


def _darken_color(hex_color: str, factor: float) -> str:
    """Darken a hex color."""
    hex_color = hex_color.lstrip('#')
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    r = int(r * (1 - factor))
    g = int(g * (1 - factor))
    b = int(b * (1 - factor))
    return f"#{r:02x}{g:02x}{b:02x}"


def _get_logo_src(path_or_url: Optional[str]) -> Optional[str]:
    """Convert a local file path or URL to a src attribute value for img tags.

    For local files, returns a base64 data URL.
    For URLs (http/https), returns the URL as-is.
    Returns None if path is None or file doesn't exist.
    """
    if not path_or_url:
        return None

    # If it's a URL, return as-is
    if path_or_url.startswith(('http://', 'https://')):
        return path_or_url

    # For local files, convert to base64 data URL
    try:
        # Handle relative paths from app directory
        if not os.path.isabs(path_or_url):
            app_dir = os.path.dirname(os.path.abspath(__file__))
            path_or_url = os.path.join(app_dir, path_or_url)

        if not os.path.exists(path_or_url):
            return None

        with open(path_or_url, 'rb') as f:
            data = f.read()

        # Determine MIME type from extension
        ext = os.path.splitext(path_or_url)[1].lower()
        mime_types = {
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.gif': 'image/gif',
            '.svg': 'image/svg+xml',
            '.ico': 'image/x-icon',
        }
        mime_type = mime_types.get(ext, 'image/png')

        b64 = base64.b64encode(data).decode('utf-8')
        return f"data:{mime_type};base64,{b64}"
    except Exception:
        return None


def graph_enabled() -> bool:
    return get_graph_config() is not None


def get_openai_client():
    api_key = get_secret("openai_api_key") or get_secret("OPENAI_API_KEY")

    if not api_key:
        st.warning("OpenAI API key not found. Calendar parsing disabled.")
        return None

    if OpenAI is not None:
        return OpenAI(api_key=api_key)

    if openai_legacy is not None:
        openai_legacy.api_key = api_key
        return openai_legacy

    st.warning("OpenAI SDK not available. Calendar parsing disabled.")
    return None

# ----------------------------
# PDF / image parsing helpers (existing behavior)
# ----------------------------
def image_to_base64(image: Image.Image) -> str:
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")


def parse_slots_from_image(image: Image.Image) -> List[Dict[str, str]]:
    """
    Use OpenAI vision to parse free/busy calendar images into slots.
    Expected JSON format:
    [
      {"date": "2025-12-03", "start": "09:00", "end": "09:30"},
      ...
    ]
    """
    client = get_openai_client()
    if not client:
        return []

    prompt = (
        "You are extracting FREE time slots from a calendar screenshot.\n"
        "Return ONLY valid JSON (no markdown) as a list of objects with keys:\n"
        "  - date (YYYY-MM-DD)\n"
        "  - start (HH:MM in 24-hour format)\n"
        "  - end (HH:MM in 24-hour format)\n"
        "  - inferred_tz (timezone abbreviation if visible, e.g. 'PST', 'EST', 'GMT', or null if not visible)\n\n"
        "Look for timezone indicators in:\n"
        "  - Calendar headers or footers\n"
        "  - Corner labels (e.g., 'Times shown in PST')\n"
        "  - Time displays with timezone suffix (e.g., '2:00 PM EST')\n"
        "  - UTC offset indicators (e.g., 'GMT-8')\n\n"
        "Only include free slots that are at least 30 minutes.\n"
        "If no slots found, return an empty list []."
    )

    b64 = image_to_base64(image)

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a helpful assistant that returns strict JSON."},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
                    ],
                },
            ],
        )
        content = resp.choices[0].message.content.strip() if resp.choices else ""
        # Strip code fences if present
        if content.startswith("```"):
            content = content.strip("`")
            if "\n" in content:
                content = content.split("\n", 1)[1].strip()

        slots = json.loads(content) if content else []
        valid_slots = []
        for s in slots:
            if isinstance(s, dict) and all(k in s for k in ("date", "start", "end")):
                slot_data = {
                    "date": str(s["date"]),
                    "start": str(s["start"]),
                    "end": str(s["end"]),
                }
                # Preserve inferred timezone if present
                if s.get("inferred_tz"):
                    slot_data["inferred_tz"] = str(s["inferred_tz"])
                valid_slots.append(slot_data)
        return valid_slots
    except json.JSONDecodeError as e:
        st.error(f"OpenAI returned invalid JSON: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI JSON parse error: {e}",
            action="parse_slots_openai",
            error_type="json_decode_error",
        )
        return []
    except Exception as e:
        st.error(f"Failed to parse availability via OpenAI: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI vision API error: {e}",
            action="parse_slots_openai",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return []


def pdf_to_images(pdf_bytes: bytes, max_pages: int = 3) -> List[Image.Image]:
    """Convert PDF to images. Returns empty list on error instead of crashing."""
    images: List[Image.Image] = []
    doc = None
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for i in range(min(len(doc), max_pages)):
            try:
                page = doc.load_page(i)
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                images.append(img)
            except Exception as e:
                log_structured(
                    LogLevel.WARNING,
                    f"Failed to process PDF page {i}: {e}",
                    action="pdf_page_process",
                    error_type="pdf_error",
                )
    except Exception as e:
        st.error(f"Failed to open PDF: {e}")
        log_structured(
            LogLevel.ERROR,
            f"Failed to open PDF: {e}",
            action="pdf_open",
            error_type="pdf_error",
            exc_info=True,
        )
    finally:
        if doc:
            doc.close()
    return images


def docx_to_text(docx_bytes: bytes) -> str:
    """
    Extract text from a Word document including paragraphs and tables.
    Returns empty string on error instead of crashing.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        st.warning("python-docx not installed. Word document parsing unavailable.")
        log_structured(
            LogLevel.ERROR,
            "python-docx not installed",
            action="docx_import",
            error_type="import_error",
        )
        return ""

    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))
        text_parts: List[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                text_parts.append(para_text)

        # Extract tables (important for calendar/availability data)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)
    except Exception as e:
        st.error(f"Failed to read Word document: {e}")
        log_structured(
            LogLevel.ERROR,
            f"Failed to read Word document: {e}",
            action="docx_read",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return ""


def docx_extract_images(docx_bytes: bytes, max_images: int = 5) -> List[Image.Image]:
    """
    Extract embedded images from a Word document.
    Returns empty list on error instead of crashing.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return []

    images: List[Image.Image] = []
    try:
        doc = DocxDocument(io.BytesIO(docx_bytes))

        # Access the document's related parts to find images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_data)).convert("RGB")
                    images.append(img)
                    if len(images) >= max_images:
                        break
                except Exception as e:
                    log_structured(
                        LogLevel.WARNING,
                        f"Failed to extract image from docx: {e}",
                        action="docx_image_extract",
                        error_type="image_error",
                    )
                    continue

        return images
    except Exception as e:
        log_structured(
            LogLevel.WARNING,
            f"Failed to extract images from Word document: {e}",
            action="docx_image_extract",
            error_type=type(e).__name__,
        )
        return []


def parse_slots_from_text(text: str) -> List[Dict[str, str]]:
    """
    Use OpenAI to parse free/busy text into slots.
    Expected JSON format:
    [
      {"date": "2025-12-03", "start": "09:00", "end": "09:30"},
      ...
    ]
    """
    if not text or not text.strip():
        return []

    client = get_openai_client()
    if not client:
        return []

    # Get current year for inference
    current_year = datetime.now().year

    prompt = f"""You are extracting FREE/AVAILABLE time slots from text describing someone's availability.

IMPORTANT RULES:
1. Only extract slots explicitly marked as FREE, AVAILABLE, or OPEN
2. Do NOT include busy/blocked/unavailable times
3. Convert all dates to YYYY-MM-DD format
4. Convert all times to 24-hour HH:MM format
5. If year is not specified, assume {current_year}
6. If end time is not specified, assume 1 hour duration
7. Only include slots that are at least 30 minutes

DATE FORMAT EXAMPLES:
- "Monday Dec 3" -> "{current_year}-12-03"
- "12/03/2025" -> "2025-12-03"
- "3rd December" -> "{current_year}-12-03"
- "Dec 3, 2025" -> "2025-12-03"

TIME FORMAT EXAMPLES:
- "9am-10am" -> start: "09:00", end: "10:00"
- "09:00-10:00" -> start: "09:00", end: "10:00"
- "9:00 AM to 10:00 AM" -> start: "09:00", end: "10:00"
- "2pm-3:30pm" -> start: "14:00", end: "15:30"

Return ONLY valid JSON as a list of objects with keys: date, start, end.
If no free slots found, return an empty list [].

TEXT TO PARSE:
{text}"""

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a helpful assistant that returns strict JSON. Never include markdown formatting."},
                {"role": "user", "content": prompt},
            ],
        )
        content = resp.choices[0].message.content.strip() if resp.choices else ""

        # Strip code fences if present (same pattern as parse_slots_from_image)
        if content.startswith("```"):
            content = content.strip("`")
            if "\n" in content:
                content = content.split("\n", 1)[1].strip()

        slots = json.loads(content) if content else []
        valid_slots = []
        for s in slots:
            if isinstance(s, dict) and all(k in s for k in ("date", "start", "end")):
                valid_slots.append({
                    "date": str(s["date"]),
                    "start": str(s["start"]),
                    "end": str(s["end"])
                })
        return valid_slots
    except json.JSONDecodeError as e:
        st.error(f"OpenAI returned invalid JSON: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI JSON parse error: {e}",
            action="parse_slots_text_openai",
            error_type="json_decode_error",
        )
        return []
    except Exception as e:
        st.error(f"Failed to parse availability via OpenAI: {e}")
        log_structured(
            LogLevel.ERROR,
            f"OpenAI text API error: {e}",
            action="parse_slots_text_openai",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return []


def ensure_session_state() -> None:
    defaults = {
        "slots": [],
        "last_graph_event_id": "",
        "last_teams_join_url": "",
        "last_invite_uid": "",
        "last_invite_ics_bytes": b"",
        "selected_timezone": get_default_timezone(),
        "candidate_timezone": get_default_timezone(),
        "duration_minutes": 30,
        # Panel interview support
        "panel_interviewers": [],  # List of {id, name, email, file, slots, timezone}
        "next_interviewer_id": 1,  # Auto-increment for unique widget keys
        "slot_filter_mode": "all_available",  # "all_available" | "any_n" | "show_all"
        "slot_filter_min_n": 1,  # Minimum N for "any_n" mode
        "computed_intersections": [],  # Intersection slots with availability metadata
        "editing_slot_index": None,  # Track which slot is being edited: (interviewer_idx, slot_idx) or None
        # Interview management UI state
        "cancelling_interview_id": None,  # ID of interview being cancelled (for confirmation dialog)
        "rescheduling_interview_id": None,  # ID of interview being rescheduled (for confirmation dialog)
        "viewing_interview_history": None,  # Event ID for viewing history
        "interview_status_filter": "All",  # Status filter for interviews list
        # Branding customization (overrides secrets) - loaded from persistent storage
        "custom_company_name": None,  # Override company name from secrets
        "custom_logo_data": None,  # Base64 encoded logo data
        "custom_primary_color": None,  # Override primary brand color
        "custom_background_color": None,  # Override background color
        "_branding_loaded": False,  # Track if branding was loaded from file
        # Audit log view state
        "audit_view_mode": "Table",  # "Timeline" | "Table" | "Raw"
        "audit_entry_limit": 300,  # Entry limit selector
        "audit_action_filter": "All",  # Action type filter
        "audit_status_filter": "All",  # Status filter (All/Success/Failed)
        "audit_search": "",  # Search term
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # Load persistent branding settings on first run
    if not st.session_state.get("_branding_loaded"):
        saved = _load_branding_settings()
        if saved:
            st.session_state["custom_company_name"] = saved.get("company_name")
            st.session_state["custom_logo_data"] = saved.get("logo_data")
            st.session_state["custom_primary_color"] = saved.get("primary_color")
            st.session_state["custom_background_color"] = saved.get("background_color")
        st.session_state["_branding_loaded"] = True


def format_slot_label(slot: Dict[str, str]) -> str:
    return f"{slot['date']} {slot['start']}–{slot['end']}"


def _merge_slots(manual_slots: List[Dict], uploaded_slots: List[Dict]) -> List[Dict]:
    """Merge slots, preferring manual over uploaded for duplicates."""
    seen = {}
    for s in manual_slots:
        key = (s["date"], s["start"], s["end"])
        seen[key] = s
    for s in uploaded_slots:
        key = (s["date"], s["start"], s["end"])
        if key not in seen:
            seen[key] = s
    return list(seen.values())


def _add_manual_slot(interviewer_idx: int, slot_date, start_time, end_time) -> bool:
    """Add a manually entered slot with validation. Returns True if successful."""
    from datetime import date as date_type, time as time_type

    errors = []

    # Validate end time is after start time
    if end_time <= start_time:
        errors.append("End time must be after start time")

    # Validate minimum duration (30 minutes)
    start_dt = datetime.combine(date.today(), start_time)
    end_dt = datetime.combine(date.today(), end_time)
    duration_minutes = (end_dt - start_dt).seconds // 60
    if duration_minutes < 30:
        errors.append("Slot must be at least 30 minutes")

    # Validate not in the past
    if slot_date < date.today():
        errors.append("Cannot add slots in the past")

    if errors:
        for err in errors:
            st.error(err)
        return False

    # Create slot in standard format
    new_slot = {
        "date": slot_date.strftime("%Y-%m-%d"),
        "start": start_time.strftime("%H:%M"),
        "end": end_time.strftime("%H:%M"),
        "source": "manual",
    }

    # Get existing slots for this interviewer
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        st.error("Invalid interviewer index")
        return False

    existing_slots = interviewers[interviewer_idx].get("slots", [])

    # Check for duplicates
    slot_key = (new_slot["date"], new_slot["start"], new_slot["end"])
    for s in existing_slots:
        if (s["date"], s["start"], s["end"]) == slot_key:
            st.warning("This slot already exists")
            return False

    existing_slots.append(new_slot)
    st.session_state["panel_interviewers"][interviewer_idx]["slots"] = existing_slots
    st.success(f"Added slot: {format_slot_label(new_slot)}")
    return True


def _delete_interviewer_slot(interviewer_idx: int, slot_idx: int) -> None:
    """Delete a slot by index from an interviewer's slots."""
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        return

    slots = interviewers[interviewer_idx].get("slots", [])
    if 0 <= slot_idx < len(slots):
        deleted = slots.pop(slot_idx)
        st.session_state["panel_interviewers"][interviewer_idx]["slots"] = slots
        st.toast(f"Deleted: {format_slot_label(deleted)}")
        st.rerun()


def _render_interviewer_slots(interviewer_idx: int, interviewer_id: int) -> None:
    """Render editable list of current slots for an interviewer."""
    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        return

    slots = interviewers[interviewer_idx].get("slots", [])

    if not slots:
        st.info("No slots added yet. Use the form above or upload a calendar.")
        return

    st.markdown(f"**Current Slots ({len(slots)}):**")

    for idx, slot in enumerate(slots):
        col_label, col_edit, col_delete = st.columns([4, 1, 1])

        with col_label:
            source_badge = " manual" if slot.get("source") == "manual" else " uploaded"
            st.text(f"{source_badge} {format_slot_label(slot)}")

        with col_edit:
            if st.button("Edit", key=f"edit_slot_{interviewer_id}_{idx}"):
                st.session_state["editing_slot_index"] = (interviewer_idx, idx)
                st.rerun()

        with col_delete:
            if st.button("Del", key=f"del_slot_{interviewer_id}_{idx}"):
                _delete_interviewer_slot(interviewer_idx, idx)

    # Clear all button
    if len(slots) > 1:
        if st.button("Clear All Slots", key=f"clear_all_{interviewer_id}", type="secondary"):
            st.session_state["panel_interviewers"][interviewer_idx]["slots"] = []
            st.rerun()


def _render_manual_slot_form(interviewer_idx: int, interviewer_id: int) -> None:
    """Render the form to add a new manual slot."""
    st.caption("Add availability slots manually instead of uploading a calendar")

    col_date, col_start, col_end, col_btn = st.columns([2, 1.5, 1.5, 1])

    with col_date:
        slot_date = st.date_input(
            "Date",
            value=date.today(),
            key=f"manual_slot_date_{interviewer_id}",
            min_value=date.today(),
        )
    with col_start:
        slot_start = st.time_input(
            "Start",
            value=time(9, 0),
            key=f"manual_slot_start_{interviewer_id}",
        )
    with col_end:
        slot_end = st.time_input(
            "End",
            value=time(10, 0),
            key=f"manual_slot_end_{interviewer_id}",
        )
    with col_btn:
        st.write("")  # Vertical spacing
        if st.button("+ Add", key=f"add_manual_slot_{interviewer_id}", type="primary"):
            if _add_manual_slot(interviewer_idx, slot_date, slot_start, slot_end):
                st.rerun()


def _render_edit_slot_form(interviewer_idx: int, interviewer_id: int) -> None:
    """Render edit form when a slot is being edited."""
    edit_info = st.session_state.get("editing_slot_index")
    if edit_info is None:
        return

    edit_interviewer_idx, edit_slot_idx = edit_info

    # Only render if this is the interviewer being edited
    if edit_interviewer_idx != interviewer_idx:
        return

    interviewers = st.session_state.get("panel_interviewers", [])
    if interviewer_idx >= len(interviewers):
        st.session_state["editing_slot_index"] = None
        return

    slots = interviewers[interviewer_idx].get("slots", [])
    if edit_slot_idx >= len(slots):
        st.session_state["editing_slot_index"] = None
        return

    slot = slots[edit_slot_idx]

    st.markdown("---")
    st.markdown(f"**Editing:** {format_slot_label(slot)}")

    col_date, col_start, col_end = st.columns(3)
    with col_date:
        new_date = st.date_input(
            "Date",
            value=datetime.strptime(slot["date"], "%Y-%m-%d").date(),
            key=f"edit_slot_date_{interviewer_id}",
        )
    with col_start:
        new_start = st.time_input(
            "Start",
            value=datetime.strptime(slot["start"], "%H:%M").time(),
            key=f"edit_slot_start_{interviewer_id}",
        )
    with col_end:
        new_end = st.time_input(
            "End",
            value=datetime.strptime(slot["end"], "%H:%M").time(),
            key=f"edit_slot_end_{interviewer_id}",
        )

    col_save, col_cancel = st.columns(2)
    with col_save:
        if st.button("Save Changes", type="primary", key=f"save_edit_{interviewer_id}"):
            # Validate
            if new_end <= new_start:
                st.error("End time must be after start time")
            elif new_date < date.today():
                st.error("Cannot set date in the past")
            else:
                duration = (datetime.combine(date.today(), new_end) - datetime.combine(date.today(), new_start)).seconds // 60
                if duration < 30:
                    st.error("Slot must be at least 30 minutes")
                else:
                    # Update the slot
                    slots[edit_slot_idx] = {
                        "date": new_date.strftime("%Y-%m-%d"),
                        "start": new_start.strftime("%H:%M"),
                        "end": new_end.strftime("%H:%M"),
                        "source": slot.get("source", "manual"),
                    }
                    st.session_state["panel_interviewers"][interviewer_idx]["slots"] = slots
                    st.session_state["editing_slot_index"] = None
                    st.success("Slot updated!")
                    st.rerun()

    with col_cancel:
        if st.button("Cancel", key=f"cancel_edit_{interviewer_id}"):
            st.session_state["editing_slot_index"] = None
            st.rerun()

    st.markdown("---")


def extract_common_timezone(slots: List[Dict[str, str]]) -> Optional[str]:
    """
    Extract the most common inferred timezone from parsed slots.

    Returns IANA timezone or None if no timezone was inferred.
    """
    from collections import Counter
    from timezone_utils import infer_timezone_from_abbreviation

    tz_abbrevs = [s.get("inferred_tz") for s in slots if s.get("inferred_tz")]
    if not tz_abbrevs:
        return None

    # Get most common abbreviation
    most_common = Counter(tz_abbrevs).most_common(1)[0][0]

    # Convert to IANA timezone name
    iana_tz, matched, _ = infer_timezone_from_abbreviation(most_common)
    return iana_tz if matched else None


# ----------------------------
# Email helpers (existing, with updated secret key names)
# ----------------------------
def build_scheduling_email(role_title: str, recruiter_name: str, slots: List[Dict[str, str]]) -> str:
    slot_lines = "\n".join([f"- {format_slot_label(s)}" for s in slots]) if slots else "- (No slots extracted)"
    return f"""Hi there,

Thanks for your time. Please choose one of the following interview times for the role: {role_title}

Available slots:
{slot_lines}

Reply with your preferred option and we will confirm the invite.

Best regards,
{recruiter_name}
Talent Acquisition
"""


def _build_logo_html(company: CompanyConfig) -> str:
    """Build logo HTML section, or empty string if no logo URL configured."""
    if not company.logo_url:
        return ""
    return f'''
    <tr>
        <td align="center" style="padding: 20px 0 10px 0;">
            <img src="{company.logo_url}" alt="{company.name}"
                 style="max-height: 60px; max-width: 200px; height: auto; display: block;" />
        </td>
    </tr>
    '''


# System font stack for cross-platform email rendering
_EMAIL_FONT_STACK = "-apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif"


def build_branded_email_html(
    candidate_name: str,
    role_title: str,
    slots: List[Dict[str, str]],
    company: CompanyConfig,
    custom_message: Optional[str] = None,
) -> str:
    """
    Build professional HTML email with company branding.

    Uses inline CSS only for email client compatibility.
    Max width 600px, table-based layout.
    """
    # Logo section
    logo_html = _build_logo_html(company)

    # Greeting with candidate name
    greeting = f"Dear {candidate_name}," if candidate_name else "Hello,"

    # Optional custom message
    custom_section = f'<p style="margin: 16px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">{custom_message}</p>' if custom_message else ""

    # Build slot list HTML
    slot_items = ""
    for slot in slots:
        slot_items += f'''
        <tr>
            <td style="padding: 10px 12px; border-left: 3px solid {company.primary_color};
                       background-color: #f8f9fa; font-family: {_EMAIL_FONT_STACK}; font-size: 15px;">
                {format_slot_label(slot)}
            </td>
        </tr>
        <tr><td style="height: 8px;"></td></tr>
        '''
    if not slots:
        slot_items = f'<tr><td style="padding: 10px 12px; color: #666; font-family: {_EMAIL_FONT_STACK};">(No slots available)</td></tr>'

    # Website link for footer
    website_link = ""
    if company.website:
        website_link = f'<p style="margin: 8px 0 0 0; font-size: 13px;"><a href="{company.website}" style="color: {company.primary_color}; text-decoration: none;">{company.website}</a></p>'

    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin: 0; padding: 0; background-color: #f4f4f4;">
    <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="background-color: #f4f4f4;">
        <tr>
            <td align="center" style="padding: 20px 10px;">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0"
                       style="max-width: 600px; width: 100%; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    {logo_html}
                    <tr>
                        <td style="padding: 32px 40px;">
                            <p style="margin: 0 0 16px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                {greeting}
                            </p>
                            <p style="margin: 0 0 16px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                Thank you for your interest in the <strong style="color: #333333;">{role_title}</strong> position at {company.name}.
                            </p>
                            {custom_section}
                            <p style="margin: 0 0 8px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                Please select one of the following available interview times:
                            </p>
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="margin: 16px 0;">
                                {slot_items}
                            </table>
                            <p style="margin: 16px 0 0 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                Simply reply to this email with your preferred time slot, and we will send you a calendar invitation with all the details.
                            </p>
                        </td>
                    </tr>
                    <!-- Signature -->
                    <tr>
                        <td style="padding: 20px 40px; background-color: #f8f9fa; border-top: 1px solid #e9ecef; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0; color: #666666; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                                Best regards,<br/>
                                <strong style="color: {company.primary_color};">{company.signature_name}</strong>
                            </p>
                            {website_link}
                        </td>
                    </tr>
                </table>
                <!-- Footer -->
                <table role="presentation" cellspacing="0" cellpadding="0" border="0" style="max-width: 600px; width: 100%;">
                    <tr>
                        <td style="padding: 16px 0; text-align: center;">
                            <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 12px; color: #999999;">
                                This email was sent from {company.sender_email}
                            </p>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''


def build_confirmation_email_html(
    candidate_name: str,
    role_title: str,
    interview_time: str,
    teams_url: Optional[str],
    interviewer_names: List[str],
    company: CompanyConfig,
) -> str:
    """
    Build confirmation email after candidate selects a slot.

    Args:
        candidate_name: Name of the candidate
        role_title: Job title/role
        interview_time: Formatted interview time in candidate's timezone
        teams_url: Optional Microsoft Teams meeting URL
        interviewer_names: List of interviewer names
        company: Company branding configuration
    """
    logo_html = _build_logo_html(company)
    greeting = f"Dear {candidate_name}," if candidate_name else "Hello,"

    # Teams meeting section
    meeting_section = ""
    if teams_url:
        meeting_section = f'''
        <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
               style="margin: 16px 0; background-color: #f0f7ff; border-left: 4px solid {company.primary_color}; border-radius: 4px;">
            <tr>
                <td style="padding: 16px;">
                    <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; font-weight: 600; color: #333333;">
                        Microsoft Teams Meeting
                    </p>
                    <a href="{teams_url}" style="color: {company.primary_color}; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; word-break: break-all;">
                        Join Meeting
                    </a>
                </td>
            </tr>
        </table>
        '''

    # Interviewers list
    interviewers_html = ""
    if interviewer_names:
        names_list = ", ".join(interviewer_names)
        interviewers_html = f'<p style="margin: 8px 0 0 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #555555;"><strong>Interviewer(s):</strong> {names_list}</p>'

    # Website link
    website_link = ""
    if company.website:
        website_link = f'<p style="margin: 8px 0 0 0; font-size: 13px;"><a href="{company.website}" style="color: {company.primary_color}; text-decoration: none;">{company.website}</a></p>'

    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin: 0; padding: 0; background-color: #f4f4f4;">
    <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="background-color: #f4f4f4;">
        <tr>
            <td align="center" style="padding: 20px 10px;">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0"
                       style="max-width: 600px; width: 100%; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    {logo_html}
                    <tr>
                        <td style="padding: 32px 40px;">
                            <p style="margin: 0 0 16px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                {greeting}
                            </p>
                            <p style="margin: 0 0 16px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                Your interview for the <strong style="color: #333333;">{role_title}</strong> position at {company.name} has been confirmed.
                            </p>
                            <!-- Interview details box -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #f9f9f9; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #333333;">
                                            <strong>Date & Time:</strong> {interview_time}
                                        </p>
                                        {interviewers_html}
                                    </td>
                                </tr>
                            </table>
                            {meeting_section}
                            <p style="margin: 16px 0 0 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                A calendar invitation has been sent to your email. If you need to reschedule, please reply to this email.
                            </p>
                        </td>
                    </tr>
                    <!-- Signature -->
                    <tr>
                        <td style="padding: 20px 40px; background-color: #f8f9fa; border-top: 1px solid #e9ecef; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0; color: #666666; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                                Best regards,<br/>
                                <strong style="color: {company.primary_color};">{company.signature_name}</strong>
                            </p>
                            {website_link}
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''


def build_branded_email_plain(
    candidate_name: str,
    role_title: str,
    slots: List[Dict[str, str]],
    company: CompanyConfig,
) -> str:
    """Plain text version of branded email for fallback."""
    greeting = f"Dear {candidate_name}," if candidate_name else "Hello,"
    slot_lines = "\n".join([f"  - {format_slot_label(s)}" for s in slots]) if slots else "  - (No slots available)"

    footer_parts = [company.signature_name]
    if company.website:
        footer_parts.append(company.website)

    return f"""{greeting}

Thank you for your interest in the {role_title} position at {company.name}.

Please select one of the following available interview times:

{slot_lines}

Simply reply to this email with your preferred time slot, and we will send you a calendar invitation.

Best regards,
{chr(10).join(footer_parts)}

---
Sent from {company.sender_email}
"""


def build_cancellation_email_html(
    candidate_name: str,
    role_title: str,
    interview_time: str,
    reason: str,
    custom_message: Optional[str],
    company: CompanyConfig,
) -> str:
    """
    Build HTML email for interview cancellation notification.

    Args:
        candidate_name: Name of the candidate
        role_title: Job title/role
        interview_time: Formatted interview time
        reason: Cancellation reason
        custom_message: Optional additional message
        company: Company branding configuration
    """
    logo_html = _build_logo_html(company)
    greeting = f"Dear {candidate_name}," if candidate_name else "Hello,"

    custom_section = ""
    if custom_message:
        custom_section = f'''
        <p style="margin: 16px 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
            {custom_message}
        </p>
        '''

    website_link = ""
    if company.website:
        website_link = f'<p style="margin: 8px 0 0 0; font-size: 13px;"><a href="{company.website}" style="color: {company.primary_color}; text-decoration: none;">{company.website}</a></p>'

    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin: 0; padding: 0; background-color: #f4f4f4;">
    <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="background-color: #f4f4f4;">
        <tr>
            <td align="center" style="padding: 20px 10px;">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0"
                       style="max-width: 600px; width: 100%; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    {logo_html}
                    <tr>
                        <td style="padding: 32px 40px;">
                            <p style="margin: 0 0 16px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                {greeting}
                            </p>
                            <!-- Cancellation notice box -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #fff3cd; border-left: 4px solid #ffc107; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 16px; font-weight: 600; color: #856404;">
                                            Interview Cancelled
                                        </p>
                                        <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #856404;">
                                            We regret to inform you that your interview for the <strong>{role_title}</strong> position
                                            scheduled for <strong>{interview_time}</strong> has been cancelled.
                                        </p>
                                    </td>
                                </tr>
                            </table>
                            <!-- Reason -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #f9f9f9; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #333333;">
                                            <strong>Reason:</strong> {reason}
                                        </p>
                                    </td>
                                </tr>
                            </table>
                            {custom_section}
                            <p style="margin: 16px 0 0 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                We apologize for any inconvenience this may cause. If you have any questions, please reply to this email.
                            </p>
                        </td>
                    </tr>
                    <!-- Signature -->
                    <tr>
                        <td style="padding: 20px 40px; background-color: #f8f9fa; border-top: 1px solid #e9ecef; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0; color: #666666; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                                Best regards,<br/>
                                <strong style="color: {company.primary_color};">{company.signature_name}</strong>
                            </p>
                            {website_link}
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''


def build_reschedule_email_html(
    candidate_name: str,
    role_title: str,
    old_time: str,
    new_time: str,
    teams_url: Optional[str],
    company: CompanyConfig,
) -> str:
    """
    Build HTML email for interview reschedule notification.

    Args:
        candidate_name: Name of the candidate
        role_title: Job title/role
        old_time: Previous interview time (formatted)
        new_time: New interview time (formatted)
        teams_url: Optional Microsoft Teams meeting URL
        company: Company branding configuration
    """
    logo_html = _build_logo_html(company)
    greeting = f"Dear {candidate_name}," if candidate_name else "Hello,"

    meeting_section = ""
    if teams_url:
        meeting_section = f'''
        <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
               style="margin: 16px 0; background-color: #f0f7ff; border-left: 4px solid {company.primary_color}; border-radius: 4px;">
            <tr>
                <td style="padding: 16px;">
                    <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; font-weight: 600; color: #333333;">
                        Microsoft Teams Meeting
                    </p>
                    <a href="{teams_url}" style="color: {company.primary_color}; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; word-break: break-all;">
                        Join Meeting
                    </a>
                </td>
            </tr>
        </table>
        '''

    website_link = ""
    if company.website:
        website_link = f'<p style="margin: 8px 0 0 0; font-size: 13px;"><a href="{company.website}" style="color: {company.primary_color}; text-decoration: none;">{company.website}</a></p>'

    return f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="margin: 0; padding: 0; background-color: #f4f4f4;">
    <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%" style="background-color: #f4f4f4;">
        <tr>
            <td align="center" style="padding: 20px 10px;">
                <table role="presentation" cellspacing="0" cellpadding="0" border="0"
                       style="max-width: 600px; width: 100%; background-color: #ffffff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    {logo_html}
                    <tr>
                        <td style="padding: 32px 40px;">
                            <p style="margin: 0 0 16px 0; color: #333333; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                {greeting}
                            </p>
                            <!-- Reschedule notice box -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #d4edda; border-left: 4px solid #28a745; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <p style="margin: 0 0 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 16px; font-weight: 600; color: #155724;">
                                            Interview Rescheduled
                                        </p>
                                        <p style="margin: 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; color: #155724;">
                                            Your interview for the <strong>{role_title}</strong> position has been rescheduled.
                                        </p>
                                    </td>
                                </tr>
                            </table>
                            <!-- Time comparison -->
                            <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%"
                                   style="margin: 16px 0; background-color: #f9f9f9; border-radius: 4px;">
                                <tr>
                                    <td style="padding: 16px;">
                                        <table role="presentation" cellspacing="0" cellpadding="0" border="0" width="100%">
                                            <tr>
                                                <td style="padding: 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; color: #666;">
                                                    <strong>Previous Time:</strong>
                                                </td>
                                                <td style="padding: 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; color: #999; text-decoration: line-through;">
                                                    {old_time}
                                                </td>
                                            </tr>
                                            <tr>
                                                <td style="padding: 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 14px; color: #155724;">
                                                    <strong>New Time:</strong>
                                                </td>
                                                <td style="padding: 8px 0; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; font-weight: 600; color: #155724;">
                                                    {new_time}
                                                </td>
                                            </tr>
                                        </table>
                                    </td>
                                </tr>
                            </table>
                            {meeting_section}
                            <p style="margin: 16px 0 0 0; color: #555555; font-family: {_EMAIL_FONT_STACK}; font-size: 15px; line-height: 1.6;">
                                An updated calendar invitation has been sent. Please update your calendar accordingly.
                            </p>
                        </td>
                    </tr>
                    <!-- Signature -->
                    <tr>
                        <td style="padding: 20px 40px; background-color: #f8f9fa; border-top: 1px solid #e9ecef; border-radius: 0 0 8px 8px;">
                            <p style="margin: 0; color: #666666; font-family: {_EMAIL_FONT_STACK}; font-size: 14px;">
                                Best regards,<br/>
                                <strong style="color: {company.primary_color};">{company.signature_name}</strong>
                            </p>
                            {website_link}
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>'''


def _smtp_cfg() -> Optional[Dict[str, Any]]:
    # New keys (preferred)
    host = get_secret("smtp_host")
    port = get_secret("smtp_port")
    username = get_secret("smtp_username")
    password = get_secret("smtp_password")
    smtp_from = get_secret("smtp_from")

    # Backward-compat keys (older app)
    if not host:
        host = get_secret("smtp_server")

    if not (host and username and password):
        return None

    return {
        "host": str(host),
        "port": int(port or 587),
        "username": str(username),
        "password": str(password),
        "from": str(smtp_from or username),
    }


def send_email_smtp(
    subject: str,
    body: str,
    to_emails: List[str],
    cc_emails: Optional[List[str]] = None,
    attachment: Optional[Dict[str, Any]] = None,
) -> bool:
    cfg = _smtp_cfg()
    if not cfg:
        st.warning("SMTP is not configured in secrets; email send is disabled.")
        return False

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = cfg["from"]
    msg["To"] = ", ".join([e for e in to_emails if e])
    if cc_emails:
        msg["Cc"] = ", ".join([e for e in cc_emails if e])
    msg.set_content(body)

    if attachment:
        msg.add_attachment(
            attachment["data"],
            maintype=attachment.get("maintype", "application"),
            subtype=attachment.get("subtype", "octet-stream"),
            filename=attachment.get("filename", "attachment.bin"),
        )

    try:
        with smtplib.SMTP(cfg["host"], cfg["port"]) as server:
            server.starttls()
            server.login(cfg["username"], cfg["password"])
            server.send_message(msg)
        return True
    except smtplib.SMTPAuthenticationError as e:
        st.error(f"SMTP authentication failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP authentication failed: {e}",
            action="smtp_send",
            error_type="smtp_auth_error",
        )
        return False
    except smtplib.SMTPException as e:
        st.error(f"SMTP send failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP send failed: {e}",
            action="smtp_send",
            error_type="smtp_error",
        )
        return False
    except Exception as e:
        st.error(f"SMTP send failed: {e}")
        log_structured(
            LogLevel.ERROR,
            f"SMTP send failed: {e}",
            action="smtp_send",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return False


def send_email_graph(
    subject: str,
    body: str,
    to_emails: List[str],
    cc_emails: Optional[List[str]] = None,
    attachment: Optional[Dict[str, Any]] = None,
    content_type: str = "Text",
) -> bool:
    """
    Send email using Microsoft Graph API.

    Args:
        subject: Email subject line
        body: Email body (plain text or HTML)
        to_emails: List of recipient email addresses
        cc_emails: Optional list of CC recipients
        attachment: Optional attachment dict with filename, data, maintype, subtype
        content_type: "Text" for plain text, "HTML" for HTML emails
    """
    cfg = get_graph_config()
    if not cfg:
        st.warning("Graph is not configured. Add graph_tenant_id, graph_client_id, graph_client_secret, graph_scheduler_mailbox in Streamlit secrets.")
        return False

    try:
        client = GraphClient(cfg)
        graph_attachment = None
        if attachment:
            graph_attachment = {
                "name": attachment.get("filename", "attachment.bin"),
                "contentBytes": attachment.get("data"),
                "contentType": f"{attachment.get('maintype', 'application')}/{attachment.get('subtype', 'octet-stream')}",
            }
        client.send_mail(
            subject=subject,
            body=body,
            to_recipients=[e for e in to_emails if e],
            cc_recipients=[e for e in (cc_emails or []) if e] or None,
            content_type=content_type,
            attachment=graph_attachment,
        )
        return True
    except Exception as e:
        st.error(f"Graph email send failed: {e}")
        return False


def fetch_unread_emails_graph() -> Tuple[List[Dict[str, Any]], Optional[str], bool]:
    """
    Fetch unread emails from scheduler mailbox via Microsoft Graph API.
    Returns (emails, error_message, is_configured) tuple.
    - error_message is None on success
    - is_configured is False if Graph credentials are missing

    Uses the same Graph credentials as calendar operations.
    """
    cfg = get_graph_config()
    if not cfg:
        return [], None, False  # Graph not configured

    try:
        from graph_client import GraphClient
        client = GraphClient(cfg)
        messages = client.fetch_unread_messages(top=50)

        emails: List[Dict[str, Any]] = []
        for msg in messages:
            from_addr = ""
            from_data = msg.get("from", {})
            if from_data:
                email_addr = from_data.get("emailAddress", {})
                from_addr = email_addr.get("address", "")

            # Get body content (prefer text, fall back to HTML)
            body_content = msg.get("bodyPreview", "")
            body_data = msg.get("body", {})
            if body_data and body_data.get("content"):
                body_content = body_data.get("content", "")
                # Strip HTML tags if content type is HTML
                if body_data.get("contentType") == "html":
                    import re
                    body_content = re.sub(r'<[^>]+>', '', body_content)
                    body_content = body_content.strip()

            emails.append({
                "id": msg.get("id", ""),
                "from": from_addr,
                "subject": msg.get("subject", ""),
                "date": msg.get("receivedDateTime", ""),
                "body": body_content,
            })

        return emails, None, True  # Success, configured

    except GraphAuthError as e:
        log_structured(
            LogLevel.ERROR,
            f"Graph authentication failed: {e}",
            action="graph_fetch_messages",
            error_type="graph_auth_error",
        )
        return [], f"Graph authentication failed: {e}", True
    except GraphAPIError as e:
        log_structured(
            LogLevel.ERROR,
            f"Graph API error: {e}",
            action="graph_fetch_messages",
            error_type="graph_api_error",
            details={"status_code": e.status_code},
        )
        return [], f"Graph API error: {e}", True
    except Exception as e:
        log_structured(
            LogLevel.ERROR,
            f"Failed to fetch emails via Graph: {e}",
            action="graph_fetch_messages",
            error_type="graph_error",
            exc_info=True,
        )
        return [], f"Failed to fetch emails: {e}", True


def detect_slot_choice_from_text(text: str, slots: List[Dict[str, str]]) -> Optional[Dict[str, str]]:
    """
    Heuristic: find a slot label or date+time mention in a reply.
    """
    t = (text or "").lower()
    for s in slots:
        label = format_slot_label(s).lower()
        if label in t:
            return s

    # fallback: look for YYYY-MM-DD and HH:MM
    m_date = re.search(r"\b(20\d{2}-\d{2}-\d{2})\b", t)
    m_time = re.search(r"\b(\d{1,2}:\d{2})\b", t)
    if m_date and m_time:
        date = m_date.group(1)
        start = m_time.group(1).zfill(5)
        for s in slots:
            if s["date"] == date and s["start"] == start:
                return s
        return {"date": date, "start": start, "end": ""}

    return None


# ----------------------------
# Graph + ICS helpers
# ----------------------------
def _make_graph_client() -> Optional[GraphClient]:
    cfg = get_graph_config()
    if not cfg:
        return None
    return GraphClient(cfg)


def _graph_event_payload(
    *,
    subject: str,
    body_html: str,
    start_local: datetime,
    end_local: datetime,
    time_zone: str,
    attendees: List[Tuple[str, str]],
    is_teams: bool,
    location: str,
) -> Dict[str, Any]:
    payload: Dict[str, Any] = {
        "subject": subject,
        "body": {"contentType": "HTML", "content": body_html},
        "start": {"dateTime": start_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": time_zone},
        "end": {"dateTime": end_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": time_zone},
        "attendees": [{"emailAddress": {"address": e, "name": n or e}, "type": "required"} for (e, n) in attendees],
    }

    if is_teams:
        payload["isOnlineMeeting"] = True
        payload["onlineMeetingProvider"] = "teamsForBusiness"
        payload["location"] = {"displayName": "Microsoft Teams"}
    else:
        payload["location"] = {"displayName": location or "Interview"}

    return payload


def _build_ics(
    *,
    organizer_email: str,
    organizer_name: str,
    attendee_emails: List[str],
    summary: str,
    description: str,
    dtstart_utc: datetime,
    dtend_utc: datetime,
    location: str,
    url: str,
    uid_hint: str,
    display_timezone: str = "UTC",
) -> bytes:
    uid = stable_uid(uid_hint, organizer_email, ",".join(attendee_emails), dtstart_utc.isoformat())
    inv = ICSInvite(
        uid=uid,
        dtstart_utc=dtstart_utc,
        dtend_utc=dtend_utc,
        summary=summary,
        description=description,
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=attendee_emails,
        location=location,
        url=url,
        display_timezone=display_timezone,
    )
    return inv.to_ics()


# ----------------------------
# Streamlit UI - Branding Components
# ----------------------------

def _apply_brand_theme(company: CompanyConfig, background_color: Optional[str] = None) -> None:
    """Apply client's brand colors to UI elements via CSS."""
    primary = company.primary_color
    primary_light = _lighten_color(primary, 0.9)
    primary_dark = _darken_color(primary, 0.2)

    # Background color CSS (only if custom color is set)
    bg_css = ""
    if background_color:
        bg_css = f"""
/* Custom background color */
.stApp, [data-testid="stAppViewContainer"] {{
    background-color: {background_color} !important;
}}
.stMain, [data-testid="stMain"], .main .block-container {{
    background-color: {background_color} !important;
}}
"""

    css = f"""<style>
{bg_css}
/* Primary buttons */
.stButton > button[kind="primary"], .stButton > button[data-testid="baseButton-primary"] {{
    background-color: {primary} !important;
    border-color: {primary} !important;
}}
.stButton > button[kind="primary"]:hover, .stButton > button[data-testid="baseButton-primary"]:hover {{
    background-color: {primary_dark} !important;
    border-color: {primary_dark} !important;
}}
/* All buttons hover effect */
.stButton > button:hover {{
    border-color: {primary} !important;
}}
/* Selected tabs */
.stTabs [data-baseweb="tab-list"] button[aria-selected="true"] {{
    border-bottom-color: {primary} !important;
    color: {primary} !important;
}}
/* Tab highlight bar */
.stTabs [data-baseweb="tab-highlight"] {{
    background-color: {primary} !important;
}}
/* Links */
a {{ color: {primary}; }}
a:hover {{ color: {primary_dark}; }}
/* Progress bars */
.stProgress > div > div > div {{
    background-color: {primary} !important;
}}
/* Selectbox/multiselect highlight */
[data-baseweb="select"] [aria-selected="true"], [data-baseweb="menu"] [aria-selected="true"] {{
    background-color: {primary_light} !important;
}}
/* Checkbox and radio when checked */
.stCheckbox [data-testid="stCheckbox"] input:checked + div {{
    background-color: {primary} !important;
    border-color: {primary} !important;
}}
/* Slider */
.stSlider [data-testid="stThumbValue"], .stSlider [data-baseweb="slider"] div[role="slider"] {{
    background-color: {primary} !important;
}}
/* Sidebar accent */
[data-testid="stSidebar"] {{
    border-right: 3px solid {primary};
}}
/* Custom branded section class */
.branded-section {{
    border-left: 4px solid {primary};
    padding-left: 16px;
    margin: 16px 0;
}}
</style>"""
    st.markdown(css, unsafe_allow_html=True)


def _render_header_full(company: CompanyConfig) -> None:
    """Render full header with logo and powered-by badge."""
    css = """<style>
.branded-header { display: flex; align-items: center; justify-content: space-between; padding: 1rem 0; border-bottom: 2px solid #f0f0f0; margin-bottom: 1.5rem; }
.client-branding { display: flex; align-items: center; gap: 16px; }
.client-logo { max-height: 50px; max-width: 180px; object-fit: contain; }
.app-title { font-size: 1.5rem; font-weight: 600; color: #333; margin: 0; }
.powered-by { display: flex; align-items: center; gap: 8px; font-size: 0.75rem; color: #888; }
.powerdash-logo { height: 20px; opacity: 0.7; }
</style>"""
    st.markdown(css, unsafe_allow_html=True)

    client_logo_html = ""
    client_logo_src = _get_logo_src(company.logo_url)
    if client_logo_src:
        client_logo_html = f'<img src="{client_logo_src}" class="client-logo" alt="{company.name}" />'

    powerdash_logo_path = get_secret("powerdash_logo_url", "logo.png")
    powerdash_logo_src = _get_logo_src(powerdash_logo_path)
    layout = get_layout_config()

    powered_by_html = ""
    if layout.show_powered_by and powerdash_logo_src:
        powered_by_html = f'<div class="powered-by"><span>Powered by</span><img src="{powerdash_logo_src}" class="powerdash-logo" alt="PowerDash" /></div>'

    header_html = f'<div class="branded-header"><div class="client-branding">{client_logo_html}<h1 class="app-title">{company.name} Interview Scheduler</h1></div>{powered_by_html}</div>'

    st.markdown(header_html, unsafe_allow_html=True)


def _render_header_compact(company: CompanyConfig) -> None:
    """Render compact single-line header."""
    logo_html = ""
    client_logo_src = _get_logo_src(company.logo_url)
    if client_logo_src:
        logo_html = f'<img src="{client_logo_src}" style="height: 32px; margin-right: 12px;" />'

    layout = get_layout_config()
    powered_by_text = '<span style="font-size: 0.7rem; color: #999;">Powered by PowerDash</span>' if layout.show_powered_by else ''

    html = f'<div style="display: flex; align-items: center; justify-content: space-between; padding: 8px 0; border-bottom: 1px solid #eee; margin-bottom: 16px;"><div style="display: flex; align-items: center;">{logo_html}<span style="font-size: 1.1rem; font-weight: 600;">{company.name} Scheduler</span></div>{powered_by_text}</div>'
    st.markdown(html, unsafe_allow_html=True)


def _render_header_minimal(company: CompanyConfig) -> None:
    """Render minimal text-only header."""
    html = f'<div style="padding: 4px 0; margin-bottom: 12px;"><span style="font-size: 1rem; color: #333;">{company.name}</span><span style="font-size: 0.8rem; color: #999; margin-left: 8px;">Interview Scheduler</span></div>'
    st.markdown(html, unsafe_allow_html=True)


def _render_branded_header(company: CompanyConfig) -> None:
    """Render header based on configured style."""
    layout = get_layout_config()

    if layout.header_style == "compact":
        _render_header_compact(company)
    elif layout.header_style == "minimal":
        _render_header_minimal(company)
    else:
        _render_header_full(company)


def _render_footer() -> None:
    """Render footer with PowerDash branding and links."""
    st.markdown("---")

    css = """<style>
.app-footer { display: flex; justify-content: space-between; align-items: center; padding: 1rem 0; color: #888; font-size: 0.8rem; }
.footer-left { display: flex; align-items: center; gap: 8px; }
.footer-logo { height: 18px; opacity: 0.6; }
.footer-links a { color: #888; text-decoration: none; margin-left: 16px; }
.footer-links a:hover { color: #555; }
</style>"""
    st.markdown(css, unsafe_allow_html=True)

    powerdash_logo_path = get_secret("powerdash_logo_url", "logo.png")
    powerdash_logo_src = _get_logo_src(powerdash_logo_path)
    current_year = datetime.now().year

    logo_html = f'<img src="{powerdash_logo_src}" class="footer-logo" alt="PowerDash" />' if powerdash_logo_src else ''
    footer_html = f'<div class="app-footer"><div class="footer-left">{logo_html}<span>&copy; {current_year} PowerDash HR. All rights reserved.</span></div><div class="footer-links"><a href="https://powerdashhr.com/support" target="_blank">Support</a><a href="https://powerdashhr.com/privacy" target="_blank">Privacy</a></div></div>'

    st.markdown(footer_html, unsafe_allow_html=True)


def _save_current_branding() -> None:
    """Save current branding settings from session state to persistent storage."""
    settings = {
        "company_name": st.session_state.get("custom_company_name"),
        "logo_data": st.session_state.get("custom_logo_data"),
        "primary_color": st.session_state.get("custom_primary_color"),
        "background_color": st.session_state.get("custom_background_color"),
    }
    # Only save if at least one setting is customized
    if any(v is not None for v in settings.values()):
        _save_branding_settings(settings)
    else:
        # Remove the file if all settings are default
        path = _get_branding_settings_path()
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass


def _render_branding_sidebar() -> None:
    """Render sidebar with branding customization settings."""
    with st.sidebar:
        st.markdown("### Settings")

        # Company name customization
        default_name = get_secret("company_name", "PowerDash HR")
        current_name = st.session_state.get("custom_company_name") or default_name

        new_name = st.text_input(
            "Company Name",
            value=current_name,
            key="branding_name_input",
            help="Customize the company name displayed in the header"
        )

        if new_name != current_name:
            if new_name and new_name != default_name:
                st.session_state["custom_company_name"] = new_name
            elif new_name == default_name:
                st.session_state["custom_company_name"] = None
            _save_current_branding()
            st.rerun()

        st.markdown("---")

        # Logo upload
        st.markdown("**Company Logo**")

        # Show current logo if set
        current_logo = st.session_state.get("custom_logo_data")
        if current_logo:
            st.image(current_logo, width=150)
            if st.button("Remove Logo", key="remove_logo_btn"):
                st.session_state["custom_logo_data"] = None
                _save_current_branding()
                st.rerun()
        else:
            default_logo_path = get_secret("company_logo_url")
            if default_logo_path:
                logo_src = _get_logo_src(default_logo_path)
                if logo_src:
                    st.image(logo_src, width=150)
                    st.caption("Default logo from settings")

        uploaded_logo = st.file_uploader(
            "Upload New Logo",
            type=["png", "jpg", "jpeg", "gif", "svg"],
            key="logo_uploader",
            help="Upload a company logo (PNG, JPG, GIF, or SVG)"
        )

        if uploaded_logo is not None:
            # Convert to base64 data URL
            data = uploaded_logo.read()
            ext = os.path.splitext(uploaded_logo.name)[1].lower()
            mime_types = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.svg': 'image/svg+xml',
            }
            mime_type = mime_types.get(ext, 'image/png')
            b64 = base64.b64encode(data).decode('utf-8')
            data_url = f"data:{mime_type};base64,{b64}"

            st.session_state["custom_logo_data"] = data_url
            _save_current_branding()
            st.rerun()

        st.markdown("---")

        # Brand color customization
        st.markdown("**Brand Color**")

        default_color = get_secret("company_primary_color", "#0066CC")
        current_color = st.session_state.get("custom_primary_color") or default_color

        # Show color preview with computed variants
        col1, col2 = st.columns([1, 2])
        with col1:
            new_color = st.color_picker(
                "Primary",
                value=current_color,
                key="brand_color_picker",
                help="Main brand color for buttons, links, and accents"
            )
        with col2:
            light_color = _lighten_color(current_color, 0.9)
            dark_color = _darken_color(current_color, 0.2)
            st.markdown(f'<div style="display:flex;gap:4px;margin-top:26px;"><div style="width:24px;height:24px;background:{current_color};border-radius:4px;" title="Primary"></div><div style="width:24px;height:24px;background:{light_color};border-radius:4px;" title="Light"></div><div style="width:24px;height:24px;background:{dark_color};border-radius:4px;" title="Dark"></div></div>', unsafe_allow_html=True)
            st.caption("Primary · Light · Dark")

        if new_color != current_color:
            if new_color != default_color:
                st.session_state["custom_primary_color"] = new_color
            else:
                st.session_state["custom_primary_color"] = None
            _save_current_branding()
            st.rerun()

        # Background color
        st.markdown("**Background Color**")
        current_bg = st.session_state.get("custom_background_color")

        col1, col2 = st.columns([1, 2])
        with col1:
            # Default to white if not set
            new_bg = st.color_picker(
                "Background",
                value=current_bg or "#FFFFFF",
                key="bg_color_picker",
                help="Page background color"
            )
        with col2:
            if current_bg:
                st.markdown(f'<div style="margin-top:26px;padding:8px;background:{current_bg};border:1px solid #ddd;border-radius:4px;font-size:11px;color:#666;">Custom</div>', unsafe_allow_html=True)
            else:
                st.markdown('<div style="margin-top:26px;padding:8px;background:#f0f0f0;border-radius:4px;font-size:11px;color:#888;">Default</div>', unsafe_allow_html=True)

        # Only save if changed from current (and not white which is default)
        if current_bg and new_bg != current_bg:
            if new_bg.upper() == "#FFFFFF":
                st.session_state["custom_background_color"] = None
            else:
                st.session_state["custom_background_color"] = new_bg
            _save_current_branding()
            st.rerun()
        elif not current_bg and new_bg.upper() != "#FFFFFF":
            st.session_state["custom_background_color"] = new_bg
            _save_current_branding()
            st.rerun()

        st.markdown("---")

        # Reset to defaults button
        if st.button("Reset to Defaults", key="reset_branding_btn"):
            st.session_state["custom_company_name"] = None
            st.session_state["custom_logo_data"] = None
            st.session_state["custom_primary_color"] = None
            st.session_state["custom_background_color"] = None
            _save_current_branding()
            st.rerun()

        # PowerDash branding at bottom
        st.markdown("---")
        powerdash_logo_src = _get_logo_src(get_secret("powerdash_logo_url", "logo.png"))
        if powerdash_logo_src:
            st.image(powerdash_logo_src, width=100)
        st.caption("Powered by PowerDash HR")


# ----------------------------
# Streamlit UI - Main App
# ----------------------------

def main() -> None:
    # Page config must come first - use secrets for initial title
    base_name = get_secret("company_name", "PowerDash HR")
    st.set_page_config(
        page_title=f"{base_name} Interview Scheduler",
        page_icon=get_secret("company_favicon_url", "🗓️"),
        layout="wide",
        initial_sidebar_state="collapsed",
    )

    ensure_session_state()

    # Render branding settings sidebar
    _render_branding_sidebar()

    # Now get company config with any session state overrides
    company = get_company_config()
    layout = get_layout_config()
    background_color = st.session_state.get("custom_background_color")

    audit = AuditLog(get_audit_log_path())

    # Apply brand theme CSS and render header
    _apply_brand_theme(company, background_color)
    _render_branded_header(company)

    tab_new, tab_inbox, tab_invites, tab_audit, tab_diag = st.tabs([
        "📝 New Request",
        "📥 Inbox",
        "📅 Interviews",
        "📜 Audit Log",
        "🔧 Diagnostics",
    ])

    # ========= TAB: New Scheduling Request =========
    with tab_new:
        st.subheader("New Scheduling Request")

        col_left, col_center, col_right = st.columns([1.2, 1.5, 1.2], gap="large")

        with col_left:
            st.markdown("#### Hiring Manager & Recruiter")
            role_title = st.text_input("Role Title", key="role_title")
            hiring_manager_name = st.text_input("Hiring Manager Name", key="hm_name")
            hiring_manager_email = st.text_input("Hiring Manager Email (required)", key="hm_email")
            recruiter_name = st.text_input("Recruiter Name", key="rec_name")
            recruiter_email = st.text_input("Recruiter Email (optional attendee)", key="rec_email")
            scheduler_mailbox = get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com")
            st.text_input("Recruiter / Scheduling Mailbox Email", value=str(scheduler_mailbox), disabled=True)

        with col_center:
            st.markdown("#### Interviewer Availability")

            # Ensure at least one interviewer exists
            if not st.session_state.get("panel_interviewers"):
                new_id = st.session_state["next_interviewer_id"]
                st.session_state["next_interviewer_id"] = new_id + 1
                st.session_state["panel_interviewers"] = [{
                    "id": new_id,
                    "name": "",
                    "email": "",
                    "file": None,
                    "slots": [],
                    "timezone": st.session_state["selected_timezone"],
                }]

            interviewers = st.session_state["panel_interviewers"]

            # Render each interviewer
            for idx, interviewer in enumerate(interviewers):
                with st.container(border=True):
                    cols = st.columns([3, 3, 1])
                    with cols[0]:
                        name = st.text_input(
                            "Name",
                            value=interviewer.get("name", ""),
                            key=f"interviewer_name_{interviewer['id']}",
                            placeholder="e.g., John Smith"
                        )
                        interviewers[idx]["name"] = name
                    with cols[1]:
                        email = st.text_input(
                            "Email",
                            value=interviewer.get("email", ""),
                            key=f"interviewer_email_{interviewer['id']}",
                            placeholder="john@company.com"
                        )
                        interviewers[idx]["email"] = email
                    with cols[2]:
                        st.write("")  # Spacing
                        # Remove button (disabled if only 1 interviewer)
                        if len(interviewers) > 1:
                            if st.button("Remove", key=f"remove_{interviewer['id']}"):
                                st.session_state["panel_interviewers"] = [
                                    i for i in interviewers if i["id"] != interviewer["id"]
                                ]
                                st.rerun()

                    # File uploader
                    uploaded = st.file_uploader(
                        f"Calendar ({interviewer.get('name') or f'Interviewer {idx+1}'})",
                        type=["pdf", "png", "jpg", "jpeg", "docx"],
                        key=f"file_{interviewer['id']}",
                    )
                    interviewers[idx]["file"] = uploaded

                    # Show slot count with breakdown
                    slot_count = len(interviewer.get("slots", []))
                    manual_count = len([s for s in interviewer.get("slots", []) if s.get("source") == "manual"])
                    uploaded_count = slot_count - manual_count
                    if slot_count > 0:
                        if manual_count > 0 and uploaded_count > 0:
                            st.caption(f"{slot_count} slot(s) ({manual_count} manual, {uploaded_count} uploaded)")
                        elif manual_count > 0:
                            st.caption(f"{slot_count} manual slot(s)")
                        else:
                            st.caption(f"{slot_count} uploaded slot(s)")

                    # Manual slot entry expander
                    with st.expander("Manual Slot Entry", expanded=False):
                        _render_manual_slot_form(idx, interviewer["id"])
                        _render_edit_slot_form(idx, interviewer["id"])
                        _render_interviewer_slots(idx, interviewer["id"])

            st.session_state["panel_interviewers"] = interviewers

            # Add interviewer button
            if st.button("+ Add Interviewer", key="add_interviewer_btn"):
                new_id = st.session_state["next_interviewer_id"]
                st.session_state["next_interviewer_id"] = new_id + 1
                st.session_state["panel_interviewers"].append({
                    "id": new_id,
                    "name": "",
                    "email": "",
                    "file": None,
                    "slots": [],
                    "timezone": st.session_state["selected_timezone"],
                })
                st.rerun()

            st.markdown("---")

            st.session_state["duration_minutes"] = st.number_input(
                "Interview duration (minutes)", min_value=15, max_value=240, step=15, value=int(st.session_state["duration_minutes"])
            )
            # Ensure selected_timezone is valid before widget renders
            if st.session_state["selected_timezone"] not in _common_timezones():
                st.session_state["selected_timezone"] = get_default_timezone()
            tz_name = st.selectbox(
                "Display timezone",
                options=_common_timezones(),
                key="selected_timezone",
            )

            # Real-time clock showing current time in selected timezone vs system timezone
            from timezone_utils import from_utc
            now_utc = datetime.now(timezone.utc)
            now_system = datetime.now().astimezone()  # System local time
            system_tz_name = now_system.strftime("%Z")  # e.g., "PST", "GMT"

            try:
                now_selected = from_utc(now_utc, tz_name)
                selected_time = now_selected.strftime("%I:%M %p %Z")
                system_time = now_system.strftime("%I:%M %p %Z")

                st.caption(f"**{tz_name}**: {selected_time} | **Your system ({system_tz_name})**: {system_time}")
            except:
                pass

            parse_btn = st.button("Parse All Availability", type="primary")

            if parse_btn:
                _parse_all_panel_availability()

            st.markdown("#### Available Time Slots")

            intersections = st.session_state.get("computed_intersections", [])
            panel_interviewers = st.session_state.get("panel_interviewers", [])
            interviewer_count = len([i for i in panel_interviewers if i.get("slots")])

            if st.session_state["slots"]:
                # Filter mode selector (only show if multiple interviewers)
                if interviewer_count > 1:
                    from slot_intersection import filter_slots_by_availability

                    filter_col1, filter_col2 = st.columns([2, 1])
                    with filter_col1:
                        filter_options = [
                            ("all_available", f"All {interviewer_count} must be available"),
                            ("any_n", "At least N are available"),
                            ("show_all", "Show all slots"),
                        ]
                        filter_mode = st.selectbox(
                            "Show slots where:",
                            options=filter_options,
                            format_func=lambda x: x[1],
                            key="slot_filter_mode_select"
                        )
                        st.session_state["slot_filter_mode"] = filter_mode[0]

                    with filter_col2:
                        if filter_mode[0] == "any_n":
                            min_n = st.number_input(
                                "Minimum N",
                                min_value=1,
                                max_value=interviewer_count,
                                value=max(1, interviewer_count - 1),
                                key="slot_filter_min_n_input"
                            )
                            st.session_state["slot_filter_min_n"] = min_n

                    # Apply filter
                    filtered_slots = filter_slots_by_availability(
                        intersections,
                        st.session_state.get("slot_filter_mode", "all_available"),
                        st.session_state.get("slot_filter_min_n", 1),
                        interviewer_count
                    )
                else:
                    filtered_slots = st.session_state["slots"]

                if not filtered_slots:
                    st.warning("No slots match the current filter. Try relaxing the availability requirement.")
                    selected_slot = None
                else:
                    st.info("Select a slot to create an invite, or generate a candidate email.")

                    # Build slot labels with availability info
                    from slot_intersection import format_slot_label_with_availability

                    def get_slot_label(slot):
                        if interviewer_count > 1:
                            return format_slot_label_with_availability(slot, interviewer_count)
                        return format_slot_label(slot)

                    slot_labels = [get_slot_label(s) for s in filtered_slots]
                    selected_label = st.selectbox("Select slot", options=slot_labels, key="selected_slot_label")
                    selected_slot = filtered_slots[slot_labels.index(selected_label)]

                    # Show availability indicator for panel interviews
                    if interviewer_count > 1 and selected_slot:
                        avail = selected_slot.get("available_count", interviewer_count)
                        total = selected_slot.get("total_interviewers", interviewer_count)
                        available_names = selected_slot.get("available_names", [])

                        if avail == total:
                            st.success(f"All {total} interviewers available")
                        elif avail >= total * 0.75:
                            missing = [
                                i.get("name") or i.get("email")
                                for i in panel_interviewers
                                if i["id"] not in selected_slot.get("available_interviewers", [])
                                and i.get("slots")
                            ]
                            st.info(f"{avail}/{total} available. Missing: {', '.join(missing) if missing else 'None'}")
                        else:
                            st.warning(f"Only {avail}/{total} interviewers available: {', '.join(available_names)}")

                    # Real-time timezone conversion preview
                    if selected_slot:
                        from timezone_utils import safe_zoneinfo, from_utc, format_time_for_display
                        try:
                            # Parse the slot time as display timezone
                            slot_dt_naive = datetime.strptime(
                                f"{selected_slot['date']}T{selected_slot['start']}:00",
                                "%Y-%m-%dT%H:%M:%S"
                            )
                            zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
                            slot_dt_local = slot_dt_naive.replace(tzinfo=zi)

                            # Convert to UTC for reference
                            from timezone_utils import to_utc
                            slot_utc = to_utc(slot_dt_local)

                            # Show conversion to common timezones
                            st.markdown("**Time Conversion Preview:**")
                            preview_tzs = ["UTC", "America/New_York", "America/Los_Angeles", "Europe/London", "Asia/Tokyo"]
                            # Add display timezone if not in list
                            if tz_name not in preview_tzs:
                                preview_tzs.insert(0, tz_name)

                            conversion_items = []
                            for preview_tz in preview_tzs:
                                try:
                                    converted = from_utc(slot_utc, preview_tz)
                                    time_str = converted.strftime("%a %b %d, %I:%M %p %Z")
                                    # Highlight the display timezone
                                    if preview_tz == tz_name:
                                        conversion_items.append(f"**{preview_tz}**: {time_str} *(selected)*")
                                    else:
                                        conversion_items.append(f"{preview_tz}: {time_str}")
                                except Exception:
                                    pass

                            st.caption(" | ".join(conversion_items[:4]))  # Show top 4
                        except (ValueError, TypeError):
                            pass  # Skip preview on invalid date

                    # DST Warning Check
                    if selected_slot:
                        from timezone_utils import is_near_dst_transition
                        try:
                            slot_date = datetime.strptime(selected_slot["date"], "%Y-%m-%d").date()
                            slot_dt = datetime.combine(slot_date, datetime.min.time())

                            # Check display timezone for DST transition
                            is_near, trans_date, trans_type = is_near_dst_transition(slot_dt, tz_name, days_threshold=7)
                            if is_near and trans_date:
                                direction = "spring forward" if trans_type == "spring_forward" else "fall back"
                                st.warning(
                                    f"DST Alert: Clocks {direction} on {trans_date.strftime('%B %d, %Y')} "
                                    f"in {tz_name}. Please verify the scheduled time."
                                )
                        except (ValueError, TypeError):
                            pass  # Skip DST check on invalid date
            else:
                st.info("No slots extracted yet. Upload availability and click Parse All Availability.")
                selected_slot = None

        with col_right:
            st.markdown("#### Candidates")
            st.caption("Enter one or more emails separated by semicolons. Format: email or Name <email>")

            candidate_input = st.text_area(
                "Candidate Email(s) (required)",
                key="multi_cand_input",
                height=80,
                placeholder="john@example.com; Jane Doe <jane@example.com>; bob@example.com"
            )

            # Parse and validate candidates
            candidate_results = parse_candidate_emails(candidate_input)
            valid_candidates = [r for r in candidate_results if r.is_valid]
            invalid_count = len(candidate_results) - len(valid_candidates)

            # Display validation results
            if candidate_results:
                if valid_candidates and not invalid_count:
                    st.success(f"All {len(valid_candidates)} candidate(s) validated")
                elif valid_candidates and invalid_count:
                    st.warning(f"{len(valid_candidates)} valid, {invalid_count} invalid candidate(s)")
                elif invalid_count:
                    st.error(f"All {invalid_count} candidate(s) have validation errors")

                # Show validation details in expander
                if len(candidate_results) > 1 or invalid_count > 0:
                    with st.expander("Validation Details", expanded=bool(invalid_count)):
                        for r in candidate_results:
                            if r.is_valid:
                                display_name = f"{r.name} ({r.email})" if r.name else r.email
                                st.markdown(f":white_check_mark: {display_name}")
                            else:
                                st.markdown(f":x: {r.original} - {r.error}")

            # Option to proceed with valid only when there are errors
            proceed_with_valid = True
            if invalid_count > 0 and len(valid_candidates) > 0:
                proceed_with_valid = st.checkbox(
                    f"Proceed with {len(valid_candidates)} valid candidate(s) only",
                    value=True,
                    key="proceed_with_valid"
                )
                if not proceed_with_valid:
                    st.info("Fix invalid candidates before proceeding, or check the box above to skip them.")

            # Scheduling mode selection (only show if multiple valid candidates)
            scheduling_mode = "individual"
            if len(valid_candidates) > 1:
                st.markdown("##### Scheduling Mode")
                scheduling_mode = st.radio(
                    "How to schedule these candidates:",
                    options=["individual", "group"],
                    format_func=lambda x: {
                        "individual": f"Individual Interviews - {len(valid_candidates)} separate invites",
                        "group": "Group Interview - All candidates in one meeting"
                    }[x],
                    key="scheduling_mode",
                    horizontal=True
                )
                if scheduling_mode == "individual":
                    st.caption("Each candidate will receive their own calendar invite.")
                else:
                    st.caption("All candidates will be invited to a single shared meeting.")

            # For backward compatibility, set candidate_email from first valid candidate
            candidate_email = valid_candidates[0].email if valid_candidates else ""
            candidate_name = valid_candidates[0].name if valid_candidates else ""

            # Candidate timezone - pre-populate with inferred timezone from calendar
            inferred_tz = extract_common_timezone(st.session_state.get("slots", []))
            if inferred_tz:
                # Update session state if inference found a timezone
                st.session_state["candidate_timezone"] = inferred_tz

            candidate_tz_default = st.session_state.get("candidate_timezone", get_default_timezone())
            candidate_tz_idx = _common_timezones().index(candidate_tz_default) if candidate_tz_default in _common_timezones() else 0

            candidate_timezone = st.selectbox(
                "Candidate Timezone",
                options=_common_timezones(),
                index=candidate_tz_idx,
                key="candidate_timezone_select",
                help="Times in the invitation will be shown in this timezone"
            )

            if inferred_tz and inferred_tz == candidate_timezone:
                st.caption("Auto-detected from calendar screenshot")

            # Show candidate's view of the selected time
            if selected_slot:
                from timezone_utils import safe_zoneinfo, to_utc, from_utc, format_datetime_for_display
                try:
                    slot_dt_naive = datetime.strptime(
                        f"{selected_slot['date']}T{selected_slot['start']}:00",
                        "%Y-%m-%dT%H:%M:%S"
                    )
                    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
                    slot_dt_local = slot_dt_naive.replace(tzinfo=zi)
                    slot_utc = to_utc(slot_dt_local)

                    # Show what candidate will see
                    candidate_view = format_datetime_for_display(slot_utc, candidate_timezone)
                    st.success(f"Candidate will see: **{candidate_view}**")
                except (ValueError, TypeError):
                    pass

            st.markdown("#### Invite details")
            is_teams = st.selectbox("Interview type", options=["Teams", "Non-Teams"], index=0, key="interview_type") == "Teams"
            subject = st.text_input("Subject/title", value=f"Interview: {role_title}" if role_title else "Interview", key="subject")
            agenda = st.text_area("Description/agenda", value="Interview discussion.", key="agenda")
            location = st.text_input("Location (non-Teams)", value="", key="location")

            include_recruiter = st.checkbox("Include recruiter as attendee", value=False, key="include_recruiter")

            st.markdown("----")
            st.markdown("#### Actions")

            # Generate branded email to candidate
            if st.button("Generate Candidate Scheduling Email"):
                company = get_company_config()
                html_body = build_branded_email_html(
                    candidate_name=candidate_name,
                    role_title=role_title or "Position",
                    slots=st.session_state["slots"],
                    company=company,
                )
                plain_body = build_branded_email_plain(
                    candidate_name=candidate_name,
                    role_title=role_title or "Position",
                    slots=st.session_state["slots"],
                    company=company,
                )
                st.session_state["candidate_email_html"] = html_body
                st.session_state["candidate_email_plain"] = plain_body

            if st.session_state.get("candidate_email_html"):
                # Preview mode toggle
                preview_mode = st.radio(
                    "Preview mode",
                    options=["Rendered", "HTML Source", "Plain Text"],
                    horizontal=True,
                    key="email_preview_mode"
                )

                if preview_mode == "Rendered":
                    st.markdown("**Email Preview (Rendered):**")
                    import streamlit.components.v1 as components
                    components.html(st.session_state["candidate_email_html"], height=500, scrolling=True)
                elif preview_mode == "HTML Source":
                    st.code(st.session_state["candidate_email_html"], language="html")
                else:
                    st.text_area("Email preview (Plain Text)", st.session_state["candidate_email_plain"], height=300)

                company = get_company_config()
                if st.button("Send Email"):
                    ok = send_email_graph(
                        subject=f"Interview Opportunity at {company.name}: {role_title}",
                        body=st.session_state["candidate_email_html"],
                        to_emails=[candidate_email] if candidate_email else [],
                        cc_emails=[recruiter_email] if recruiter_email else None,
                        content_type="HTML",
                    )
                    audit.log(
                        "graph_sent_scheduling_email" if ok else "graph_send_failed",
                        actor=recruiter_email or "",
                        candidate_email=candidate_email or "",
                        hiring_manager_email=hiring_manager_email or "",
                        recruiter_email=recruiter_email or "",
                        role_title=role_title or "",
                        payload={"subject": f"Interview Opportunity at {company.name}: {role_title}"},
                        status="success" if ok else "failed",
                        error_message="" if ok else "Graph email send failed",
                    )
                    if ok:
                        st.success("Email sent.")
                    else:
                        st.error("Email send failed (see message above).")

            # Create Graph event
            # Collect panel interviewers from session state
            panel_interviewers_for_invite = [
                {"name": i.get("name", ""), "email": i.get("email", "")}
                for i in st.session_state.get("panel_interviewers", [])
                if i.get("email")  # Only include interviewers with valid emails
            ]

            # Determine if we have enough info to create invite
            has_interviewers = bool(panel_interviewers_for_invite) or bool(hiring_manager_email)
            has_valid_candidates = len(valid_candidates) > 0 and (proceed_with_valid or not invalid_count)
            create_disabled = not (selected_slot and has_interviewers and has_valid_candidates)

            # Button text reflects number of candidates
            button_text = "Create & Send Interview Invite"
            if len(valid_candidates) > 1:
                button_text = f"Create & Send {len(valid_candidates)} Interview Invite(s)"

            if st.button(button_text, disabled=create_disabled):
                with st.spinner(f"Scheduling {len(valid_candidates)} interview(s)..."):
                    results = _handle_multi_candidate_invite(
                        audit=audit,
                        selected_slot=selected_slot,
                        tz_name=tz_name,
                        candidate_timezone=candidate_timezone,
                        duration_minutes=int(st.session_state["duration_minutes"]),
                        role_title=role_title,
                        subject=subject,
                        agenda=agenda,
                        location=location,
                        is_teams=is_teams,
                        candidates=valid_candidates,
                        hiring_manager=(hiring_manager_email, hiring_manager_name),
                        recruiter=(recruiter_email, recruiter_name),
                        include_recruiter=include_recruiter,
                        panel_interviewers=panel_interviewers_for_invite if panel_interviewers_for_invite else None,
                        scheduling_mode=scheduling_mode,
                    )

                # Display batch results
                _render_batch_results(results)

            # ICS fallback download button (available after generation)
            if st.session_state.get("last_invite_ics_bytes"):
                st.download_button(
                    "Download .ics (Add to calendar)",
                    data=st.session_state["last_invite_ics_bytes"],
                    file_name="powerdash_interview_invite.ics",
                    mime="text/calendar",
                )
                audit.log(
                    "ics_downloaded",
                    actor=recruiter_email or "",
                    candidate_email=candidate_email or "",
                    hiring_manager_email=hiring_manager_email or "",
                    recruiter_email=recruiter_email or "",
                    role_title=role_title or "",
                    event_id=st.session_state.get("last_graph_event_id", ""),
                    payload={"uid": st.session_state.get("last_invite_uid", "")},
                    status="success",
                )

                # Optional email ICS via Graph
                if st.button("Email .ics (optional)"):
                    ok = send_email_graph(
                        subject=subject,
                        body=agenda,
                        to_emails=[candidate_email, hiring_manager_email] + ([recruiter_email] if include_recruiter and recruiter_email else []),
                        attachment={
                            "data": st.session_state["last_invite_ics_bytes"],
                            "maintype": "text",
                            "subtype": "calendar",
                            "filename": "invite.ics",
                        },
                    )
                    audit.log(
                        "graph_sent_ics" if ok else "graph_send_failed",
                        actor=recruiter_email or "",
                        candidate_email=candidate_email or "",
                        hiring_manager_email=hiring_manager_email or "",
                        recruiter_email=recruiter_email or "",
                        role_title=role_title or "",
                        event_id=st.session_state.get("last_graph_event_id", ""),
                        payload={"uid": st.session_state.get("last_invite_uid", "")},
                        status="success" if ok else "failed",
                        error_message="" if ok else "Graph email send failed",
                    )
                    st.success("ICS emailed.") if ok else st.error("Failed to email ICS.")

    # ========= TAB: Scheduler Inbox =========
    with tab_inbox:
        st.subheader("Scheduler Inbox")
        st.caption("Reads unread emails from the scheduler mailbox via Microsoft Graph API.")
        emails, graph_error, is_configured = fetch_unread_emails_graph()
        if not is_configured:
            st.warning("Graph API is not configured. Add graph_tenant_id, graph_client_id, graph_client_secret, and graph_scheduler_mailbox to your secrets.")
        elif graph_error:
            st.error(f"Failed to fetch emails: {graph_error}")
        elif not emails:
            st.success("✓ Connected to mailbox. No unread emails found.")
        else:
            st.write(f"Found {len(emails)} unread email(s).")
            for i, e in enumerate(emails, start=1):
                with st.expander(f"{i}. {e['subject'] or '(no subject)'} — {e['from']}"):
                    st.write(e.get("date", ""))
                    body = e.get("body", "")
                    st.text_area("Body", body, height=160)

                    if st.session_state["slots"]:
                        choice = detect_slot_choice_from_text(body, st.session_state["slots"])
                        if choice:
                            st.success(f"Detected slot choice: {choice.get('date')} {choice.get('start')}")
                        else:
                            st.info("No slot choice detected from this email.")

    # ========= TAB: Calendar Invites =========
    with tab_invites:
        st.subheader("Interview Management")
        st.caption("Manage scheduled interviews: reschedule, cancel, or view history.")

        def _format_candidates_display(interview_row: Dict[str, Any]) -> str:
            """Format candidate display for table, handling multi-candidate interviews."""
            candidates_json = interview_row.get("candidates_json")
            if candidates_json:
                try:
                    candidates = json.loads(candidates_json)
                    if len(candidates) > 2:
                        return f"{candidates[0]['email']} +{len(candidates)-1} more"
                    return "; ".join(c['email'] for c in candidates)
                except (json.JSONDecodeError, KeyError, TypeError):
                    pass
            return interview_row.get("candidate_email", "")

        def _format_interview_type(interview_row: Dict[str, Any]) -> str:
            """Format interview type indicator."""
            is_group = interview_row.get("is_group_interview")
            is_panel = interview_row.get("is_panel_interview")
            if is_group:
                return "Group"
            elif is_panel:
                return "Panel"
            return "Individual"

        def _get_status_badge(status: str) -> str:
            """Return status with emoji badge."""
            status_badges = {
                "pending": "Pending",
                "confirmed": "Confirmed",
                "rescheduled": "Rescheduled",
                "cancelled": "Cancelled",
                "completed": "Completed",
                "no_show": "No Show",
                "created": "Created",
                "scheduled": "Scheduled",
            }
            return status_badges.get(status.lower() if status else "", status or "Unknown")

        # --- Audit Log Rendering Functions ---

        def _render_audit_timeline(entries: List[Dict[str, Any]]):
            """Render audit log as timeline view with status indicators."""
            if not entries:
                st.info("No audit entries to display.")
                return

            # Action color mapping
            action_colors = {
                "graph_create_event": "#28a745",      # Green - success
                "graph_create_group_event": "#28a745",
                "interview_rescheduled": "#ffc107",   # Yellow - change
                "graph_reschedule_event": "#ffc107",
                "interview_cancelled": "#dc3545",     # Red - cancellation
                "graph_cancel_event": "#dc3545",
                "notification_sent": "#17a2b8",       # Blue - notification
                "candidate_notification_sent": "#17a2b8",
                "email_sent": "#17a2b8",
                "graph_sent_scheduling_email": "#17a2b8",
            }
            default_color = "#6c757d"  # Gray

            current_date = None

            for entry in entries:
                # Date separator
                timestamp = entry.get("timestamp", "")
                entry_date = timestamp[:12] if timestamp else ""

                if entry_date and entry_date != current_date:
                    current_date = entry_date
                    st.markdown(f"### {current_date}")

                # Get color for this action
                action_code = entry.get("action_code", "")
                color = action_colors.get(action_code, default_color)

                # Status icon
                status = entry.get("status", "")
                if status == "success":
                    status_icon = "✅"
                elif status == "failed":
                    status_icon = "❌"
                else:
                    status_icon = "ℹ️"

                # Entry card
                with st.container():
                    col_status, col_content = st.columns([0.5, 11.5])

                    with col_status:
                        st.markdown(f"### {status_icon}")

                    with col_content:
                        st.markdown(f"**{entry.get('action_display', '')}**")
                        st.markdown(entry.get("summary", ""))

                        # Meta line
                        time_part = timestamp[13:] if len(timestamp) > 13 else ""
                        meta = f"_{time_part} by {entry.get('actor_display', 'System')}_"
                        st.caption(meta)

                        # Details if present
                        if entry.get("details"):
                            st.caption(entry["details"])

                        # Expandable raw data
                        with st.expander("View details"):
                            st.json(entry.get("raw", {}))

                    st.markdown("---")

        def _render_audit_table(entries: List[Dict[str, Any]]):
            """Render audit log as formatted table view."""
            if not entries:
                st.info("No audit entries to display.")
                return

            # Status icon mapping
            def get_status_icon(status: str) -> str:
                if status == "success":
                    return "✅"
                elif status == "failed":
                    return "❌"
                return "ℹ️"

            table_data = [
                {
                    "Time": e.get("timestamp", ""),
                    "Status": get_status_icon(e.get("status", "")),
                    "Action": e.get("action_display", ""),
                    "Summary": e.get("summary", ""),
                    "Actor": e.get("actor_display", ""),
                }
                for e in entries
            ]

            st.dataframe(
                table_data,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Status": st.column_config.TextColumn(width="small"),
                    "Summary": st.column_config.TextColumn(width="large"),
                }
            )

            # Expandable row details
            with st.expander("View entry details"):
                idx = st.number_input(
                    "Entry index (0 = most recent)",
                    min_value=0,
                    max_value=max(0, len(entries) - 1),
                    value=0,
                    key="audit_table_row_idx"
                )
                selected = entries[int(idx)]
                st.markdown(f"**{selected.get('action_display', '')}**")
                st.markdown(selected.get("summary", ""))
                if selected.get("details"):
                    st.caption(selected["details"])
                st.json(selected.get("raw", {}))

        def _render_audit_raw(entries: List[Dict[str, Any]]):
            """Render audit log as raw dataframe (original view)."""
            if not entries:
                st.info("No audit entries to display.")
                return

            st.dataframe(
                [
                    {
                        "timestamp_utc": r.get("timestamp_utc", ""),
                        "action": r.get("action", ""),
                        "status": r.get("status", ""),
                        "candidate": r.get("candidate_email", ""),
                        "hiring_manager": r.get("hiring_manager_email", ""),
                        "event_id": r.get("event_id", ""),
                        "error": (r.get("error_message", "")[:80] + "…")
                            if r.get("error_message") and len(r.get("error_message", "")) > 80
                            else (r.get("error_message") or ""),
                    }
                    for r in entries
                ],
                use_container_width=True,
                hide_index=True,
            )

            with st.expander("Show raw payload for a row"):
                idx = st.number_input(
                    "Row index (0 = most recent)",
                    min_value=0,
                    max_value=max(0, len(entries) - 1),
                    value=0,
                    key="audit_raw_row_idx"
                )
                st.json(entries[int(idx)])

        def _render_interview_history(audit_instance: AuditLog, event_id: str):
            """Display complete history for selected interview."""
            st.markdown("#### Interview History")

            history = audit_instance.get_interview_history(event_id)

            if not history:
                st.info("No history entries for this interview.")
                return

            # Action color mapping
            action_colors = {
                "graph_create_event": "#28a745",     # Green
                "interview_rescheduled": "#ffc107", # Yellow
                "interview_cancelled": "#dc3545",   # Red
                "notification_sent": "#17a2b8",      # Blue
            }

            for entry in history:
                action = entry.get("action", "")
                timestamp = entry.get("timestamp_utc", "")[:16]
                actor = entry.get("actor", "System") or "System"
                status = entry.get("status", "")

                color = action_colors.get(action, "#6c757d")

                col_time, col_action, col_status = st.columns([2, 4, 2])
                with col_time:
                    st.caption(timestamp)
                with col_action:
                    st.markdown(f"<span style='color: {color}; font-weight: 600;'>{action}</span>", unsafe_allow_html=True)
                with col_status:
                    st.caption(f"{status} | {actor[:20]}")

                # Show details in expander
                payload_json = entry.get("payload_json")
                if payload_json:
                    with st.expander("Details", expanded=False):
                        try:
                            st.json(json.loads(payload_json))
                        except json.JSONDecodeError:
                            st.text(payload_json)

                st.markdown("---")

        # Filter controls
        col_filter1, col_filter2 = st.columns(2)
        with col_filter1:
            status_filter = st.selectbox(
                "Filter by status",
                options=["All", "Pending", "Confirmed", "Rescheduled", "Cancelled", "Scheduled", "Created"],
                key="invites_status_filter"
            )
        with col_filter2:
            search_term = st.text_input("Search candidate/role", key="invites_search", placeholder="Search...")

        # Get interviews with optional filter
        filter_value = None if status_filter == "All" else status_filter.lower()
        interviews = audit.list_interviews(limit=200, status_filter=filter_value)

        # Apply search filter
        if search_term:
            search_lower = search_term.lower()
            interviews = [
                r for r in interviews
                if search_lower in (r.get("candidate_email", "") or "").lower()
                or search_lower in (r.get("role_title", "") or "").lower()
            ]

        if not interviews:
            st.info("No interviews match the current filters. Create an invite from the first tab.")
        else:
            # Show compact table
            st.dataframe(
                [
                    {
                        "created": r["created_utc"][:10] if r.get("created_utc") else "",
                        "role": r.get("role_title", ""),
                        "type": _format_interview_type(r),
                        "candidate(s)": _format_candidates_display(r),
                        "start_utc": r.get("start_utc", "")[:16] if r.get("start_utc") else "",
                        "status": _get_status_badge(r.get("last_status", "")),
                        "event_id": (r.get("graph_event_id", "") or "")[:12] + "..." if r.get("graph_event_id") else "",
                    }
                    for r in interviews
                ],
                use_container_width=True,
                hide_index=True,
            )

            # Export controls section
            with st.expander("Export Data", expanded=False):
                col_tz, col_fields, col_status = st.columns(3)

                with col_tz:
                    export_tz = st.selectbox(
                        "Timezone",
                        options=_common_timezones(),
                        index=_tz_index(get_default_timezone()),
                        key="export_interviews_tz"
                    )

                with col_fields:
                    include_all_fields = st.checkbox(
                        "Include all fields",
                        value=False,
                        help="Include extended fields like Teams links, Event IDs",
                        key="export_all_fields"
                    )

                with col_status:
                    export_status_filter = st.multiselect(
                        "Filter by status",
                        options=["Pending", "Confirmed", "Rescheduled", "Cancelled", "Completed"],
                        default=["Pending", "Confirmed", "Rescheduled"],
                        key="export_status_filter"
                    )

                col_date_range, col_custom = st.columns(2)

                with col_date_range:
                    date_range = st.selectbox(
                        "Date range",
                        options=["All time", "Today", "This week", "This month", "Last 30 days", "Custom"],
                        key="export_date_range"
                    )

                date_from = None
                date_to = None
                if date_range == "Custom":
                    with col_custom:
                        col_from, col_to = st.columns(2)
                        with col_from:
                            date_from = st.date_input("From", key="export_date_from")
                        with col_to:
                            date_to = st.date_input("To", key="export_date_to")

                # Filter and generate export
                filtered_interviews = filter_interviews_for_export(
                    interviews,
                    status_filter=export_status_filter if export_status_filter else None,
                    date_range=date_range,
                    date_from=date_from,
                    date_to=date_to,
                )

                col_info, col_download = st.columns([3, 1])

                with col_info:
                    st.caption(f"{len(filtered_interviews)} interview(s) match the filters")

                with col_download:
                    if filtered_interviews:
                        csv_bytes = export_interviews_csv(
                            filtered_interviews,
                            display_timezone=export_tz,
                            include_all_fields=include_all_fields,
                        )
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        st.download_button(
                            label="Download CSV",
                            data=csv_bytes,
                            file_name=f"interviews_export_{timestamp}.csv",
                            mime="text/csv",
                            key="download_interviews_csv"
                        )
                    else:
                        st.button("Download CSV", disabled=True, key="download_disabled")

            st.markdown("----")

            # Check for active confirmation dialogs
            cancelling_id = st.session_state.get("cancelling_interview_id")
            rescheduling_id = st.session_state.get("rescheduling_interview_id")
            viewing_history_id = st.session_state.get("viewing_interview_history")

            # Interview History View
            if viewing_history_id:
                st.markdown("### Interview History")
                if st.button("Back to Interview List", key="back_from_history"):
                    st.session_state["viewing_interview_history"] = None
                    st.rerun()
                _render_interview_history(audit, viewing_history_id)

            # Cancellation Confirmation Dialog
            elif cancelling_id:
                cancel_row = next((r for r in interviews if r.get("graph_event_id") == cancelling_id), None)
                if cancel_row:
                    st.markdown("### Confirm Cancellation")
                    st.warning(f"""
                    **You are about to cancel this interview:**
                    - **Role:** {cancel_row.get('role_title', '')}
                    - **Candidate:** {cancel_row.get('candidate_email', '')}
                    - **Time:** {cancel_row.get('start_utc', '')}

                    This action cannot be undone. The candidate will receive a cancellation notice.
                    """)

                    cancel_reason = st.selectbox(
                        "Cancellation reason",
                        options=[
                            "Candidate requested",
                            "Position filled",
                            "Position closed",
                            "Interviewer unavailable",
                            "Scheduling conflict",
                            "Other"
                        ],
                        key="cancel_reason_select"
                    )

                    custom_reason = ""
                    if cancel_reason == "Other":
                        custom_reason = st.text_input("Please specify reason", key="cancel_custom_reason")

                    notify_candidate = st.checkbox("Send cancellation email to candidate", value=True, key="cancel_notify")
                    candidate_message = ""
                    if notify_candidate:
                        candidate_message = st.text_area(
                            "Message to candidate (optional)",
                            placeholder="We apologize for any inconvenience...",
                            key="cancel_message"
                        )

                    col_confirm, col_back = st.columns(2)
                    with col_confirm:
                        if st.button("Confirm Cancellation", type="primary", key="confirm_cancel_btn"):
                            final_reason = custom_reason if cancel_reason == "Other" else cancel_reason
                            _handle_cancel(
                                audit=audit,
                                event_id=cancelling_id,
                                context_row=cancel_row,
                                reason=final_reason,
                                notify_candidate=notify_candidate,
                                candidate_message=candidate_message,
                            )
                            st.session_state["cancelling_interview_id"] = None
                            st.rerun()
                    with col_back:
                        if st.button("Go Back", key="cancel_back_btn"):
                            st.session_state["cancelling_interview_id"] = None
                            st.rerun()
                else:
                    st.session_state["cancelling_interview_id"] = None
                    st.rerun()

            # Reschedule Confirmation Dialog
            elif rescheduling_id:
                resched_row = next((r for r in interviews if r.get("graph_event_id") == rescheduling_id), None)
                if resched_row:
                    st.markdown("### Reschedule Interview")

                    # Display current info
                    st.info(f"""
                    **Rescheduling interview:**
                    - **Role:** {resched_row.get('role_title', '')}
                    - **Candidate:** {resched_row.get('candidate_email', '')}
                    - **Current Time:** {resched_row.get('start_utc', '')}
                    """)

                    display_tz = st.selectbox(
                        "Timezone",
                        options=_common_timezones(),
                        index=_tz_index(resched_row.get("display_timezone")),
                        key="resched_tz"
                    )

                    try:
                        start_local = from_utc(datetime.fromisoformat(resched_row["start_utc"]), display_tz)
                    except Exception:
                        start_local = None

                    col_date, col_time = st.columns(2)
                    with col_date:
                        new_date = st.date_input(
                            "New date",
                            value=start_local.date() if start_local else datetime.now().date(),
                            key="resched_date"
                        )
                    with col_time:
                        new_time = st.time_input(
                            "New time",
                            value=start_local.time().replace(second=0, microsecond=0) if start_local else datetime.now().time().replace(second=0, microsecond=0),
                            key="resched_time"
                        )

                    new_duration = st.number_input(
                        "Duration (minutes)",
                        min_value=15,
                        max_value=240,
                        step=15,
                        value=int(resched_row.get("duration_minutes") or 30),
                        key="resched_duration"
                    )

                    reschedule_reason = st.selectbox(
                        "Reason for reschedule",
                        options=[
                            "Candidate requested",
                            "Interviewer unavailable",
                            "Scheduling conflict",
                            "Time zone adjustment",
                            "Other"
                        ],
                        key="reschedule_reason_select"
                    )

                    notify_candidate = st.checkbox(
                        "Send update notification to candidate",
                        value=True,
                        key="reschedule_notify"
                    )

                    col_confirm, col_back = st.columns(2)
                    with col_confirm:
                        if st.button("Confirm Reschedule", type="primary", key="confirm_resched_btn"):
                            _handle_reschedule(
                                audit=audit,
                                event_id=rescheduling_id,
                                new_date=new_date,
                                new_time=new_time,
                                duration_minutes=int(new_duration),
                                tz_name=display_tz,
                                context_row=resched_row,
                                reason=reschedule_reason,
                                notify_candidate=notify_candidate,
                            )
                            st.session_state["rescheduling_interview_id"] = None
                            st.rerun()
                    with col_back:
                        if st.button("Go Back", key="resched_back_btn"):
                            st.session_state["rescheduling_interview_id"] = None
                            st.rerun()
                else:
                    st.session_state["rescheduling_interview_id"] = None
                    st.rerun()

            # Default view - Interview management
            else:
                st.markdown("#### Manage Interview")

                # Only show non-cancelled interviews for management
                manageable = [r for r in interviews if r.get("last_status", "").lower() != "cancelled"]
                if not manageable:
                    st.info("No active interviews to manage. All interviews are cancelled or none exist.")
                else:
                    event_ids = [r["graph_event_id"] for r in manageable if r.get("graph_event_id")]
                    if event_ids:
                        selected_event_id = st.selectbox(
                            "Select interview",
                            options=event_ids,
                            format_func=lambda x: next(
                                (f"{r.get('role_title', '')} - {r.get('candidate_email', '')} ({r.get('start_utc', '')[:10]})"
                                 for r in manageable if r.get("graph_event_id") == x),
                                x
                            ),
                            key="manage_event_select"
                        )
                        selected_row = next((r for r in manageable if r.get("graph_event_id") == selected_event_id), None)

                        if selected_row:
                            # Show interview details
                            st.markdown(f"""
                            **Selected Interview:**
                            - **Role:** {selected_row.get('role_title', '')}
                            - **Candidate:** {selected_row.get('candidate_email', '')}
                            - **Start:** {selected_row.get('start_utc', '')}
                            - **Status:** {_get_status_badge(selected_row.get('last_status', ''))}
                            """)

                            # Action buttons
                            col_resched, col_cancel, col_history = st.columns(3)
                            with col_resched:
                                if st.button("Reschedule", type="primary", key="btn_reschedule"):
                                    st.session_state["rescheduling_interview_id"] = selected_event_id
                                    st.rerun()
                            with col_cancel:
                                if st.button("Cancel", type="secondary", key="btn_cancel"):
                                    st.session_state["cancelling_interview_id"] = selected_event_id
                                    st.rerun()
                            with col_history:
                                if st.button("View History", key="btn_history"):
                                    st.session_state["viewing_interview_history"] = selected_event_id
                                    st.rerun()

    # ========= TAB: Audit Log =========
    with tab_audit:
        st.subheader("Audit Log")
        st.caption("Complete history of all scheduling actions.")

        # Controls row
        col_view, col_limit = st.columns([3, 1])

        with col_view:
            view_mode = st.radio(
                "View",
                options=["Table", "Timeline", "Raw"],
                horizontal=True,
                key="audit_view_mode"
            )

        with col_limit:
            entry_limit = st.selectbox(
                "Entries",
                options=[100, 300, 500],
                index=1,
                key="audit_entry_limit"
            )

        # Filter controls
        col_action, col_status, col_search = st.columns([2, 1, 2])

        with col_action:
            action_options = ["All"] + sorted(set(AUDIT_ACTION_DESCRIPTIONS.values()))
            action_filter = st.selectbox(
                "Filter by action",
                options=action_options,
                key="audit_action_filter"
            )

        with col_status:
            status_filter = st.selectbox(
                "Status",
                options=["All", "Success", "Failed"],
                key="audit_status_filter"
            )

        with col_search:
            audit_search = st.text_input(
                "Search",
                placeholder="candidate, role...",
                key="audit_search"
            )

        # Fetch entries
        raw_entries = audit.list_recent_audit(limit=entry_limit)

        if not raw_entries:
            st.info("No audit entries yet.")
        else:
            # Apply filters
            filtered_entries = filter_audit_entries(
                raw_entries,
                action_filter=action_filter,
                status_filter=status_filter,
                search_term=audit_search,
            )

            if not filtered_entries:
                st.warning("No entries match the current filters.")
            else:
                # Format entries for display
                formatted_entries = [format_audit_entry_human(e) for e in filtered_entries]

                # Export button
                col_count, col_export = st.columns([3, 1])

                with col_count:
                    st.caption(f"Showing {len(formatted_entries)} of {len(raw_entries)} entries")

                with col_export:
                    csv_bytes = export_audit_log_csv(formatted_entries)
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    st.download_button(
                        label="Export CSV",
                        data=csv_bytes,
                        file_name=f"audit_log_{timestamp}.csv",
                        mime="text/csv",
                        key="download_audit_csv"
                    )

                st.markdown("---")

                # Render based on view mode
                if view_mode == "Timeline":
                    _render_audit_timeline(formatted_entries)
                elif view_mode == "Table":
                    _render_audit_table(formatted_entries)
                else:
                    _render_audit_raw(filtered_entries)

    # ========= TAB: Graph Diagnostics =========
    with tab_diag:
        st.subheader("Graph Diagnostics")
        st.caption("Use this to verify Graph auth and mailbox access. No secrets are displayed.")

        cfg = get_graph_config()
        if not cfg:
            st.warning("Graph is not configured. Add graph_tenant_id, graph_client_id, graph_client_secret, graph_scheduler_mailbox in Streamlit secrets.")
        else:
            st.code(f"Scheduler mailbox: {cfg.scheduler_mailbox}")
            client = GraphClient(cfg)

            if st.button("Test token acquisition"):
                try:
                    token = client.get_token(force_refresh=True)
                    st.success(f"Token OK (length={len(token)})")
                    audit.log("graph_token_ok", payload={"length": len(token)}, status="success")
                except Exception as e:
                    st.error(str(e))
                    audit.log("graph_token_failed", payload={"error": str(e)}, status="failed", error_message=str(e))

            if st.button("Test calendar read (top 5)"):
                try:
                    data = client.test_calendar_read(top=5)
                    st.success("Calendar read OK.")
                    st.json(data)
                    audit.log("graph_calendar_read_ok", payload={"top": 5}, status="success")
                except GraphAPIError as e:
                    st.error(f"{e} — details below")
                    st.json(e.response_json)
                    audit.log("graph_calendar_read_failed", payload=e.response_json, status="failed", error_message=str(e))

            st.markdown("----")
            st.markdown("#### Dummy event (optional)")
            dry_run = st.checkbox("Dry run (do not create)", value=True)
            tz_name = st.selectbox("Timezone", options=_common_timezones(), index=_common_timezones().index(get_default_timezone()) if get_default_timezone() in _common_timezones() else 0)
            dt = datetime.now().replace(second=0, microsecond=0) + timedelta(hours=2)
            date = st.date_input("Date", value=dt.date())
            time = st.time_input("Time", value=dt.time())
            start_local = datetime.combine(date, time).replace(tzinfo=_zoneinfo(tz_name))
            end_local = start_local + timedelta(minutes=30)

            if st.button("Create dummy event"):
                try:
                    start_dt = {"dateTime": start_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name}
                    end_dt = {"dateTime": end_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name}
                    out = client.create_dummy_event("PowerDash Diagnostics", start_dt, end_dt, dry_run=dry_run)
                    st.success("OK")
                    st.json(out)
                    audit.log("graph_dummy_event_ok", payload={"dry_run": dry_run}, status="success")
                except GraphAPIError as e:
                    st.error(f"{e}")
                    st.json(e.response_json)
                    audit.log("graph_dummy_event_failed", payload=e.response_json, status="failed", error_message=str(e))

    # Render footer if enabled
    if layout.show_footer:
        _render_footer()


# ----------------------------
# Internal UI handlers
# ----------------------------
def _parse_availability_upload(upload) -> List[Dict[str, str]]:
    data = upload.read()
    name = (upload.name or "").lower()
    slots: List[Dict[str, str]] = []

    if name.endswith(".pdf"):
        imgs = pdf_to_images(data, max_pages=3)
        for img in imgs:
            slots.extend(parse_slots_from_image(img))

    elif name.endswith(".docx"):
        # Strategy: Extract text + tables, then also check embedded images

        # 1. Parse text content (paragraphs + tables)
        text = docx_to_text(data)
        if text:
            slots.extend(parse_slots_from_text(text))

        # 2. Extract and parse embedded images (optional enhancement)
        embedded_images = docx_extract_images(data, max_images=3)
        for img in embedded_images:
            slots.extend(parse_slots_from_image(img))

    else:
        # Assume image file (png, jpg, jpeg)
        img = Image.open(io.BytesIO(data)).convert("RGB")
        slots.extend(parse_slots_from_image(img))

    # De-duplicate slots by (date, start, end) tuple
    uniq = {(s["date"], s["start"], s["end"]): s for s in slots}
    return list(uniq.values())


def _parse_all_panel_availability() -> None:
    """Parse availability for all interviewers and compute intersection.

    Handles both uploaded files and manually-entered slots.
    """
    from slot_intersection import (
        normalize_slots_to_utc,
        merge_adjacent_slots,
        compute_intersection,
    )

    interviewers = st.session_state.get("panel_interviewers", [])
    tz_name = st.session_state["selected_timezone"]
    min_duration = st.session_state["duration_minutes"]

    all_interviewer_slots: Dict[int, List] = {}
    interviewer_names: Dict[int, str] = {}
    parse_errors = []
    total_uploaded = 0
    total_manual = 0

    for interviewer in interviewers:
        # Get existing manual slots (preserve them)
        existing_manual_slots = [s for s in interviewer.get("slots", []) if s.get("source") == "manual"]
        total_manual += len(existing_manual_slots)

        try:
            if interviewer.get("file"):
                # Reset file position before reading
                interviewer["file"].seek(0)
                # Parse the uploaded file
                uploaded_slots = _parse_availability_upload(interviewer["file"])
                # Mark uploaded slots with source
                for s in uploaded_slots:
                    s["source"] = "uploaded"
                total_uploaded += len(uploaded_slots)
                # Merge manual + uploaded, preferring manual for duplicates
                interviewer["slots"] = _merge_slots(existing_manual_slots, uploaded_slots)
            elif existing_manual_slots:
                # No file but has manual slots - keep them
                interviewer["slots"] = existing_manual_slots

            # Include interviewer if they have any slots
            if interviewer.get("slots"):
                # Build interviewer name for display
                name = interviewer.get("name") or interviewer.get("email") or f"Interviewer {interviewer['id']}"
                interviewer_names[interviewer["id"]] = name

                # Normalize to UTC for intersection
                normalized = normalize_slots_to_utc(interviewer["slots"], tz_name)
                merged = merge_adjacent_slots(normalized)
                all_interviewer_slots[interviewer["id"]] = merged

        except Exception as e:
            interviewer_name = interviewer.get("name") or f"Interviewer {interviewer.get('id', '?')}"
            parse_errors.append(f"{interviewer_name}: {e}")

    if parse_errors:
        for err in parse_errors:
            st.error(err)

    # Compute intersection
    if all_interviewer_slots:
        intersections = compute_intersection(
            all_interviewer_slots,
            min_duration_minutes=min_duration,
            display_timezone=tz_name,
            interviewer_names=interviewer_names,
        )
        st.session_state["computed_intersections"] = intersections

        # Also update legacy "slots" for backward compatibility
        st.session_state["slots"] = intersections

        num_interviewers = len(all_interviewer_slots)
        total_slots = total_uploaded + total_manual

        if num_interviewers == 1:
            if total_manual > 0 and total_uploaded > 0:
                st.success(f"Processed {total_slots} slot(s) ({total_manual} manual, {total_uploaded} uploaded).")
            elif total_manual > 0:
                st.success(f"Processed {total_manual} manual slot(s).")
            else:
                st.success(f"Extracted {total_uploaded} slot(s) from uploaded file.")
        else:
            full_overlap = sum(1 for s in intersections if s.get("is_full_overlap", False))
            source_info = []
            if total_manual > 0:
                source_info.append(f"{total_manual} manual")
            if total_uploaded > 0:
                source_info.append(f"{total_uploaded} uploaded")
            source_str = f" ({', '.join(source_info)})" if source_info else ""
            st.success(
                f"Processed {total_slots} total slots{source_str} from {num_interviewers} interviewers. "
                f"Found {len(intersections)} intersection slot(s) ({full_overlap} with all available)."
            )
    else:
        st.warning("No availability found. Please upload calendars or add slots manually.")


def _zoneinfo(tz_name: str):
    """Get ZoneInfo with validation. Falls back to UTC if invalid."""
    zi, was_valid = safe_zoneinfo(tz_name, fallback="UTC")
    if not was_valid:
        st.warning(f"Invalid timezone '{tz_name}', using UTC")
    return zi


def _common_timezones() -> List[str]:
    # Keep concise; you can expand later.
    return [
        "UTC",
        "Europe/London",
        "Europe/Dublin",
        "Europe/Paris",
        "Europe/Rome",
        "Europe/Berlin",
        "America/New_York",
        "America/Chicago",
        "America/Denver",
        "America/Los_Angeles",
        "America/Toronto",
        "America/Sao_Paulo",
        "Asia/Dubai",
        "Asia/Kolkata",
        "Asia/Singapore",
        "Asia/Tokyo",
        "Australia/Sydney",
    ]


def _tz_index(tz_name: str | None) -> int:
    tzs = _common_timezones()
    if tz_name and tz_name in tzs:
        return tzs.index(tz_name)
    return tzs.index(get_default_timezone()) if get_default_timezone() in tzs else 0


def _render_batch_results(results: List[SchedulingResult]) -> None:
    """Display results of batch scheduling operation."""
    if not results:
        return

    successful = [r for r in results if r.success]
    failed = [r for r in results if not r.success]

    # Summary
    st.markdown("### Scheduling Results")

    if successful and not failed:
        st.success(f"Successfully scheduled all {len(successful)} interview(s)")
    elif successful and failed:
        st.warning(f"Scheduled {len(successful)} interview(s), {len(failed)} failed")
    else:
        st.error(f"Failed to schedule all {len(failed)} interview(s)")

    # Success details
    if successful:
        with st.expander(f"Successful ({len(successful)})", expanded=True):
            for r in successful:
                display = f"{r.candidate_name} ({r.candidate_email})" if r.candidate_name else r.candidate_email
                st.markdown(f":white_check_mark: **{display}**")
                if r.teams_url:
                    st.link_button("Open Teams Link", r.teams_url, key=f"teams_{r.event_id}_{r.candidate_email}")

    # Failure details
    if failed:
        with st.expander(f"Failed ({len(failed)})", expanded=True):
            for r in failed:
                display = f"{r.candidate_name} ({r.candidate_email})" if r.candidate_name else r.candidate_email
                st.markdown(f":x: **{display}**")
                st.caption(f"Error: {r.error}")


def _handle_multi_candidate_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidates: List[CandidateValidationResult],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
    scheduling_mode: str = "individual",
) -> List[SchedulingResult]:
    """
    Handle creating invites for multiple candidates.
    Returns list of SchedulingResult for UI display.
    """
    results: List[SchedulingResult] = []

    if scheduling_mode == "group" and len(candidates) > 1:
        # Group interview: single invite with all candidates
        result = _create_group_invite(
            audit=audit,
            selected_slot=selected_slot,
            tz_name=tz_name,
            candidate_timezone=candidate_timezone,
            duration_minutes=duration_minutes,
            role_title=role_title,
            subject=subject,
            agenda=agenda,
            location=location,
            is_teams=is_teams,
            candidates=candidates,
            hiring_manager=hiring_manager,
            recruiter=recruiter,
            include_recruiter=include_recruiter,
            panel_interviewers=panel_interviewers,
        )
        results.append(result)
    else:
        # Individual interviews: one invite per candidate
        for candidate in candidates:
            if not candidate.is_valid:
                results.append(SchedulingResult(
                    candidate_email=candidate.original,
                    candidate_name=candidate.name,
                    success=False,
                    event_id=None,
                    teams_url=None,
                    error=candidate.error or "Invalid candidate"
                ))
                continue

            result = _create_individual_invite(
                audit=audit,
                selected_slot=selected_slot,
                tz_name=tz_name,
                candidate_timezone=candidate_timezone,
                duration_minutes=duration_minutes,
                role_title=role_title,
                subject=subject,
                agenda=agenda,
                location=location,
                is_teams=is_teams,
                candidate=(candidate.email, candidate.name),
                hiring_manager=hiring_manager,
                recruiter=recruiter,
                include_recruiter=include_recruiter,
                panel_interviewers=panel_interviewers,
            )
            results.append(result)

    return results


def _create_individual_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidate: Tuple[str, str],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
) -> SchedulingResult:
    """
    Create a single individual interview invite.
    Returns SchedulingResult with success/failure status.
    """
    candidate_email_raw, candidate_name = candidate
    hm_email_raw, hm_name = hiring_manager
    rec_email_raw, rec_name = recruiter

    # === INPUT VALIDATION ===
    if not is_valid_timezone(tz_name):
        tz_name = "UTC"
    if not is_valid_timezone(candidate_timezone):
        candidate_timezone = tz_name

    try:
        candidate_email = validate_email(candidate_email_raw, "Candidate email")
        hm_email = validate_email(hm_email_raw, "Hiring manager email")
        rec_email = validate_email_optional(rec_email_raw, "Recruiter email")
    except ValidationError as e:
        return SchedulingResult(
            candidate_email=candidate_email_raw,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error=e.message
        )

    try:
        validate_slot(selected_slot)
    except ValidationError as e:
        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error=f"Invalid slot: {e.message}"
        )

    # Parse selected slot into a local datetime
    try:
        start_local_naive = datetime.fromisoformat(f"{selected_slot['date']}T{selected_slot['start']}:00")
    except ValueError as e:
        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error=f"Invalid date/time format: {e}"
        )

    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
    start_local = start_local_naive.replace(tzinfo=zi)
    end_local = start_local + timedelta(minutes=duration_minutes)
    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # Build attendees list
    attendees: List[Tuple[str, str]] = [(candidate_email, candidate_name)]
    is_panel = panel_interviewers and len(panel_interviewers) > 1
    validated_panel: List[Dict[str, str]] = []

    if panel_interviewers:
        seen_emails = {candidate_email}
        for pi in panel_interviewers:
            pi_email = (pi.get("email") or "").strip().lower()
            if pi_email and pi_email not in seen_emails:
                try:
                    validated_email = validate_email(pi_email, "Panel interviewer email")
                    validated_panel.append({"name": pi.get("name", ""), "email": validated_email})
                    attendees.append((validated_email, pi.get("name", "")))
                    seen_emails.add(validated_email)
                except ValidationError:
                    pass
    else:
        attendees.append((hm_email, hm_name))

    if include_recruiter and rec_email:
        attendees.append((rec_email, rec_name))

    organizer_email = str(get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"))
    organizer_name = "PowerDash Scheduler"

    effective_subject = subject
    if is_panel:
        if not subject.startswith("Panel Interview"):
            effective_subject = f"Panel Interview: {role_title} - {candidate_name}"

    # Generate ICS fallback
    ics_bytes = _build_ics(
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=[a[0] for a in attendees],
        summary=effective_subject,
        description=agenda,
        dtstart_utc=start_utc,
        dtend_utc=end_utc,
        location=("Microsoft Teams" if is_teams else (location or "Interview")),
        url="",
        uid_hint=f"{role_title}|{candidate_email}|{hm_email}",
        display_timezone=candidate_timezone,
    )
    st.session_state["last_invite_ics_bytes"] = ics_bytes
    st.session_state["last_invite_uid"] = stable_uid(f"{role_title}|{candidate_email}|{hm_email}", organizer_email, start_utc.isoformat())

    client = _make_graph_client()
    if not client:
        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error="Graph API not configured"
        )

    # Format time display for candidate's timezone
    from timezone_utils import format_datetime_for_display
    candidate_time_display = format_datetime_for_display(start_utc, candidate_timezone)

    body_html = f"<p><strong>Interview Time (your timezone): {candidate_time_display}</strong></p>"
    if is_panel and validated_panel:
        body_html += "<p><strong>Interview Panel:</strong></p><ul>"
        for pi in validated_panel:
            name = pi.get("name") or pi.get("email", "")
            body_html += f"<li>{name}</li>"
        body_html += "</ul>"
    if agenda:
        body_html += f"<p>{agenda.replace(chr(10), '<br>')}</p>"

    payload = _graph_event_payload(
        subject=effective_subject,
        body_html=body_html,
        start_local=start_local,
        end_local=end_local,
        time_zone=candidate_timezone,
        attendees=attendees,
        is_teams=is_teams,
        location=location,
    )

    try:
        created = client.create_event(payload)
        event_id = created.get("id", "")
        teams_url = ""
        if is_teams:
            teams_url = (created.get("onlineMeeting") or {}).get("joinUrl") or ""

        # Serialize panel interviewers for database storage
        panel_json = ""
        if validated_panel:
            panel_json = json.dumps(validated_panel)

        audit.upsert_interview(
            role_title=role_title,
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            duration_minutes=duration_minutes,
            start_utc=iso_utc(start_utc),
            end_utc=iso_utc(end_utc),
            display_timezone=tz_name,
            candidate_timezone=candidate_timezone,
            graph_event_id=event_id,
            teams_join_url=teams_url,
            subject=effective_subject,
            last_status="created",
            panel_interviewers_json=panel_json,
            is_panel_interview=is_panel,
        )

        audit.log(
            "graph_create_event",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            event_id=event_id,
            payload=payload,
            status="success",
        )

        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=True,
            event_id=event_id,
            teams_url=teams_url,
            error=None
        )

    except (GraphAuthError, GraphAPIError) as e:
        details = getattr(e, "response_json", None)
        audit.log(
            "graph_create_failed",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            payload={"error": str(e), "details": details},
            status="failed",
            error_message=str(e),
        )
        return SchedulingResult(
            candidate_email=candidate_email,
            candidate_name=candidate_name,
            success=False,
            event_id=None,
            teams_url=None,
            error=str(e)
        )


def _create_group_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidates: List[CandidateValidationResult],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
) -> SchedulingResult:
    """
    Create a single group interview invite with all candidates.
    Returns SchedulingResult with success/failure status.
    """
    hm_email_raw, hm_name = hiring_manager
    rec_email_raw, rec_name = recruiter

    # === INPUT VALIDATION ===
    if not is_valid_timezone(tz_name):
        tz_name = "UTC"
    if not is_valid_timezone(candidate_timezone):
        candidate_timezone = tz_name

    try:
        hm_email = validate_email(hm_email_raw, "Hiring manager email")
        rec_email = validate_email_optional(rec_email_raw, "Recruiter email")
    except ValidationError as e:
        return SchedulingResult(
            candidate_email=", ".join(c.email or c.original for c in candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error=e.message
        )

    try:
        validate_slot(selected_slot)
    except ValidationError as e:
        return SchedulingResult(
            candidate_email=", ".join(c.email or c.original for c in candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error=f"Invalid slot: {e.message}"
        )

    # Parse selected slot into a local datetime
    try:
        start_local_naive = datetime.fromisoformat(f"{selected_slot['date']}T{selected_slot['start']}:00")
    except ValueError as e:
        return SchedulingResult(
            candidate_email=", ".join(c.email or c.original for c in candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error=f"Invalid date/time format: {e}"
        )

    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
    start_local = start_local_naive.replace(tzinfo=zi)
    end_local = start_local + timedelta(minutes=duration_minutes)
    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # Build attendees list with all valid candidates
    valid_candidates = [c for c in candidates if c.is_valid and c.email]
    attendees: List[Tuple[str, str]] = [(c.email, c.name) for c in valid_candidates]

    is_panel = panel_interviewers and len(panel_interviewers) > 1
    validated_panel: List[Dict[str, str]] = []

    if panel_interviewers:
        seen_emails = {c.email for c in valid_candidates}
        for pi in panel_interviewers:
            pi_email = (pi.get("email") or "").strip().lower()
            if pi_email and pi_email not in seen_emails:
                try:
                    validated_email = validate_email(pi_email, "Panel interviewer email")
                    validated_panel.append({"name": pi.get("name", ""), "email": validated_email})
                    attendees.append((validated_email, pi.get("name", "")))
                    seen_emails.add(validated_email)
                except ValidationError:
                    pass
    else:
        attendees.append((hm_email, hm_name))

    if include_recruiter and rec_email:
        attendees.append((rec_email, rec_name))

    organizer_email = str(get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"))
    organizer_name = "PowerDash Scheduler"

    # Use group interview subject
    effective_subject = f"Group Interview: {role_title}"
    if subject and not subject.startswith("Interview:"):
        effective_subject = subject

    # Build candidates JSON for storage
    candidates_json = json.dumps([
        {"email": c.email, "name": c.name} for c in valid_candidates
    ])

    # Generate ICS fallback
    primary_candidate = valid_candidates[0] if valid_candidates else None
    ics_bytes = _build_ics(
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=[a[0] for a in attendees],
        summary=effective_subject,
        description=agenda,
        dtstart_utc=start_utc,
        dtend_utc=end_utc,
        location=("Microsoft Teams" if is_teams else (location or "Interview")),
        url="",
        uid_hint=f"{role_title}|group|{hm_email}",
        display_timezone=candidate_timezone,
    )
    st.session_state["last_invite_ics_bytes"] = ics_bytes
    st.session_state["last_invite_uid"] = stable_uid(f"{role_title}|group|{hm_email}", organizer_email, start_utc.isoformat())

    client = _make_graph_client()
    if not client:
        return SchedulingResult(
            candidate_email=", ".join(c.email for c in valid_candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error="Graph API not configured"
        )

    # Format time display for candidate's timezone
    from timezone_utils import format_datetime_for_display
    candidate_time_display = format_datetime_for_display(start_utc, candidate_timezone)

    body_html = f"<p><strong>Interview Time (your timezone): {candidate_time_display}</strong></p>"
    body_html += f"<p><strong>Candidates ({len(valid_candidates)}):</strong></p><ul>"
    for c in valid_candidates:
        display = c.name if c.name else c.email
        body_html += f"<li>{display}</li>"
    body_html += "</ul>"

    if is_panel and validated_panel:
        body_html += "<p><strong>Interview Panel:</strong></p><ul>"
        for pi in validated_panel:
            name = pi.get("name") or pi.get("email", "")
            body_html += f"<li>{name}</li>"
        body_html += "</ul>"
    if agenda:
        body_html += f"<p>{agenda.replace(chr(10), '<br>')}</p>"

    payload = _graph_event_payload(
        subject=effective_subject,
        body_html=body_html,
        start_local=start_local,
        end_local=end_local,
        time_zone=candidate_timezone,
        attendees=attendees,
        is_teams=is_teams,
        location=location,
    )

    try:
        created = client.create_event(payload)
        event_id = created.get("id", "")
        teams_url = ""
        if is_teams:
            teams_url = (created.get("onlineMeeting") or {}).get("joinUrl") or ""

        # Serialize panel interviewers for database storage
        panel_json = ""
        if validated_panel:
            panel_json = json.dumps(validated_panel)

        # Store with primary candidate email for backward compatibility
        audit.upsert_interview(
            role_title=role_title,
            candidate_email=primary_candidate.email if primary_candidate else "",
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            duration_minutes=duration_minutes,
            start_utc=iso_utc(start_utc),
            end_utc=iso_utc(end_utc),
            display_timezone=tz_name,
            candidate_timezone=candidate_timezone,
            graph_event_id=event_id,
            teams_join_url=teams_url,
            subject=effective_subject,
            last_status="created",
            panel_interviewers_json=panel_json,
            is_panel_interview=is_panel,
            candidates_json=candidates_json,
            is_group_interview=True,
        )

        audit.log(
            "graph_create_group_event",
            actor=rec_email or "",
            candidate_email=", ".join(c.email for c in valid_candidates),
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            event_id=event_id,
            payload={"candidates_count": len(valid_candidates)},
            status="success",
        )

        return SchedulingResult(
            candidate_email=", ".join(c.email for c in valid_candidates),
            candidate_name=f"Group ({len(valid_candidates)} candidates)",
            success=True,
            event_id=event_id,
            teams_url=teams_url,
            error=None
        )

    except (GraphAuthError, GraphAPIError) as e:
        details = getattr(e, "response_json", None)
        audit.log(
            "graph_create_group_failed",
            actor=rec_email or "",
            candidate_email=", ".join(c.email for c in valid_candidates),
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            payload={"error": str(e), "details": details, "candidates_count": len(valid_candidates)},
            status="failed",
            error_message=str(e),
        )
        return SchedulingResult(
            candidate_email=", ".join(c.email for c in valid_candidates),
            candidate_name="Group",
            success=False,
            event_id=None,
            teams_url=None,
            error=str(e)
        )


def _handle_create_invite(
    *,
    audit: AuditLog,
    selected_slot: Dict[str, str],
    tz_name: str,
    candidate_timezone: str,
    duration_minutes: int,
    role_title: str,
    subject: str,
    agenda: str,
    location: str,
    is_teams: bool,
    candidate: Tuple[str, str],
    hiring_manager: Tuple[str, str],
    recruiter: Tuple[str, str],
    include_recruiter: bool,
    panel_interviewers: Optional[List[Dict[str, str]]] = None,
) -> None:
    candidate_email_raw, candidate_name = candidate
    hm_email_raw, hm_name = hiring_manager
    rec_email_raw, rec_name = recruiter

    # === INPUT VALIDATION ===
    # Validate timezones
    if not is_valid_timezone(tz_name):
        st.warning(f"Invalid display timezone '{tz_name}', using UTC")
        tz_name = "UTC"

    if not is_valid_timezone(candidate_timezone):
        st.warning(f"Invalid candidate timezone '{candidate_timezone}', using display timezone")
        candidate_timezone = tz_name

    # Validate emails
    try:
        candidate_email = validate_email(candidate_email_raw, "Candidate email")
        hm_email = validate_email(hm_email_raw, "Hiring manager email")
        rec_email = validate_email_optional(rec_email_raw, "Recruiter email")
    except ValidationError as e:
        st.error(f"Validation error: {e.message}")
        return

    # Validate slot format
    try:
        validate_slot(selected_slot)
    except ValidationError as e:
        st.error(f"Invalid time slot: {e.message}")
        return

    # Parse selected slot into a local datetime
    try:
        start_local_naive = datetime.fromisoformat(f"{selected_slot['date']}T{selected_slot['start']}:00")
    except ValueError as e:
        st.error(f"Selected slot has invalid date/time format: {e}")
        return

    zi, _ = safe_zoneinfo(tz_name, fallback="UTC")
    start_local = start_local_naive.replace(tzinfo=zi)
    end_local = start_local + timedelta(minutes=duration_minutes)

    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # === IDEMPOTENCY CHECK ===
    existing = audit.interview_exists(
        candidate_email=candidate_email,
        hiring_manager_email=hm_email,
        role_title=role_title,
        start_utc=iso_utc(start_utc),
    )
    if existing:
        st.warning(
            f"An interview already exists for this candidate at this time. "
            f"Event ID: {existing.get('graph_event_id', 'N/A')}"
        )
        # Use a unique key based on the slot to avoid Streamlit duplicate key errors
        checkbox_key = f"force_dup_{selected_slot['date']}_{selected_slot['start']}"
        if not st.checkbox("Create duplicate anyway?", key=checkbox_key):
            return

    attendees: List[Tuple[str, str]] = [(candidate_email, candidate_name)]

    # Build attendees from panel interviewers if provided, otherwise use hiring manager
    is_panel = panel_interviewers and len(panel_interviewers) > 1
    validated_panel: List[Dict[str, str]] = []

    if panel_interviewers:
        seen_emails = {candidate_email}  # Avoid duplicating candidate
        for pi in panel_interviewers:
            pi_email = (pi.get("email") or "").strip().lower()
            if pi_email and pi_email not in seen_emails:
                try:
                    validated_email = validate_email(pi_email, "Panel interviewer email")
                    validated_panel.append({"name": pi.get("name", ""), "email": validated_email})
                    attendees.append((validated_email, pi.get("name", "")))
                    seen_emails.add(validated_email)
                except ValidationError:
                    pass  # Skip invalid emails
    else:
        # Fall back to single hiring manager (backward compatibility)
        attendees.append((hm_email, hm_name))

    if include_recruiter and rec_email:
        attendees.append((rec_email, rec_name))

    organizer_email = str(get_secret("graph_scheduler_mailbox", "scheduling@powerdashhr.com"))
    organizer_name = "PowerDash Scheduler"

    # Update subject for panel interviews
    effective_subject = subject
    if is_panel:
        if not subject.startswith("Panel Interview"):
            effective_subject = f"Panel Interview: {role_title} - {candidate_name}"

    # Always generate ICS (so we have a fallback even if Graph works)
    ics_bytes = _build_ics(
        organizer_email=organizer_email,
        organizer_name=organizer_name,
        attendee_emails=[a[0] for a in attendees],
        summary=effective_subject,
        description=agenda,
        dtstart_utc=start_utc,
        dtend_utc=end_utc,
        location=("Microsoft Teams" if is_teams else (location or "Interview")),
        url="",
        uid_hint=f"{role_title}|{candidate_email}|{hm_email}",
        display_timezone=candidate_timezone,
    )
    st.session_state["last_invite_ics_bytes"] = ics_bytes
    st.session_state["last_invite_uid"] = stable_uid(f"{role_title}|{candidate_email}|{hm_email}", organizer_email, start_utc.isoformat())
    audit.log(
        "ics_generated",
        actor=rec_email or "",
        candidate_email=candidate_email,
        hiring_manager_email=hm_email,
        recruiter_email=rec_email or "",
        role_title=role_title,
        payload={"uid": st.session_state["last_invite_uid"]},
        status="success",
    )

    client = _make_graph_client()
    if not client:
        st.warning("Graph is not configured. Using .ics fallback only.")
        return

    # Format time display for candidate's timezone
    from timezone_utils import format_datetime_for_display
    candidate_time_display = format_datetime_for_display(start_utc, candidate_timezone)

    # Build body with candidate-friendly time display
    body_html = f"<p><strong>Interview Time (your timezone): {candidate_time_display}</strong></p>"

    # Add panel members to body if panel interview
    if is_panel and validated_panel:
        body_html += "<p><strong>Interview Panel:</strong></p><ul>"
        for pi in validated_panel:
            name = pi.get("name") or pi.get("email", "")
            body_html += f"<li>{name}</li>"
        body_html += "</ul>"

    if agenda:
        body_html += f"<p>{agenda.replace(chr(10), '<br>')}</p>"

    payload = _graph_event_payload(
        subject=effective_subject,
        body_html=body_html,
        start_local=start_local,
        end_local=end_local,
        time_zone=candidate_timezone,  # Use candidate timezone for calendar event
        attendees=attendees,
        is_teams=is_teams,
        location=location,
    )

    try:
        created = client.create_event(payload)
        event_id = created.get("id", "")
        teams_url = ""
        if is_teams:
            teams_url = (created.get("onlineMeeting") or {}).get("joinUrl") or ""
        st.session_state["last_graph_event_id"] = event_id
        st.session_state["last_teams_join_url"] = teams_url

        # Re-generate ICS including Teams URL if present (better fallback)
        if teams_url:
            st.session_state["last_invite_ics_bytes"] = _build_ics(
                organizer_email=organizer_email,
                organizer_name=organizer_name,
                attendee_emails=[a[0] for a in attendees],
                summary=effective_subject,
                description=agenda,
                dtstart_utc=start_utc,
                dtend_utc=end_utc,
                location="Microsoft Teams",
                url=teams_url,
                uid_hint=f"{role_title}|{candidate_email}|{hm_email}",
                display_timezone=candidate_timezone,
            )

        audit.log(
            "graph_create_event",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            event_id=event_id,
            payload=payload,
            status="success",
        )

        # Serialize panel interviewers for database storage
        panel_json = ""
        if validated_panel:
            import json as _json
            panel_json = _json.dumps(validated_panel)

        audit.upsert_interview(
            role_title=role_title,
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            duration_minutes=duration_minutes,
            start_utc=iso_utc(start_utc),
            end_utc=iso_utc(end_utc),
            display_timezone=tz_name,
            candidate_timezone=candidate_timezone,
            graph_event_id=event_id,
            teams_join_url=teams_url,
            subject=effective_subject,
            last_status="created",
            panel_interviewers_json=panel_json,
            is_panel_interview=is_panel,
        )

        st.success("Invite created and sent via Microsoft Graph.")
        if teams_url:
            st.link_button("Open Teams meeting link", teams_url)
    except (GraphAuthError, GraphAPIError) as e:
        details = getattr(e, "response_json", None)
        st.error("Graph scheduling failed. .ics fallback is available for download.")
        if details:
            st.json(details)
        audit.log(
            "graph_create_failed",
            actor=rec_email or "",
            candidate_email=candidate_email,
            hiring_manager_email=hm_email,
            recruiter_email=rec_email or "",
            role_title=role_title,
            payload={"error": str(e), "details": details},
            status="failed",
            error_message=str(e),
        )


def _extract_candidate_name_from_context(context_row: Dict[str, Any]) -> str:
    """
    Extract candidate name from interview context row.

    Tries candidates_json first (for group interviews), then falls back to
    parsing from candidate_email if no name is found.
    """
    # Try to get name from candidates_json (for group/multi-candidate interviews)
    candidates_json = context_row.get("candidates_json")
    if candidates_json:
        try:
            candidates = json.loads(candidates_json)
            if candidates and isinstance(candidates, list) and len(candidates) > 0:
                first_candidate = candidates[0]
                if isinstance(first_candidate, dict) and first_candidate.get("name"):
                    return first_candidate["name"]
        except (json.JSONDecodeError, KeyError, TypeError):
            pass

    # Fall back to empty string - email will be used as fallback in templates
    return ""


def _format_interview_time_for_candidate(
    utc_time_str: str,
    candidate_timezone: Optional[str],
    display_timezone: Optional[str] = None,
) -> str:
    """
    Format interview time in candidate's timezone for notifications.

    Args:
        utc_time_str: ISO format UTC time string
        candidate_timezone: Candidate's preferred timezone
        display_timezone: Fallback display timezone

    Returns:
        Human-readable formatted time string
    """
    from timezone_utils import from_utc, is_valid_timezone

    # Determine which timezone to use
    tz_to_use = None
    if candidate_timezone and is_valid_timezone(candidate_timezone):
        tz_to_use = candidate_timezone
    elif display_timezone and is_valid_timezone(display_timezone):
        tz_to_use = display_timezone

    try:
        # Parse the UTC time
        dt_utc = datetime.fromisoformat(utc_time_str.replace("+00:00", "").replace("Z", ""))
        dt_utc = dt_utc.replace(tzinfo=timezone.utc)

        if tz_to_use:
            # Convert to candidate/display timezone
            dt_local = from_utc(dt_utc, tz_to_use)
            tz_abbrev = dt_local.strftime("%Z") or tz_to_use
            return dt_local.strftime(f"%A, %B %d, %Y at %I:%M %p {tz_abbrev}")
        else:
            # Fall back to UTC display
            return dt_utc.strftime("%A, %B %d, %Y at %I:%M %p UTC")
    except Exception:
        return utc_time_str


def _send_cancellation_email(
    client: GraphClient,
    candidate_email: str,
    candidate_name: str,
    role_title: str,
    interview_time: str,
    reason: str,
    custom_message: str,
    company: CompanyConfig,
) -> bool:
    """Send cancellation notification email to candidate."""
    try:
        html_body = build_cancellation_email_html(
            candidate_name=candidate_name,
            role_title=role_title,
            interview_time=interview_time,
            reason=reason,
            custom_message=custom_message if custom_message else None,
            company=company,
        )
        client.send_mail(
            subject=f"Interview Cancelled: {role_title} at {company.name}",
            body=html_body,
            to_recipients=[candidate_email],
            content_type="HTML",
        )
        log_structured(
            LogLevel.INFO,
            f"Cancellation email sent to {candidate_email}",
            action="cancellation_email_sent",
        )
        return True
    except Exception as e:
        log_structured(
            LogLevel.ERROR,
            f"Failed to send cancellation email: {e}",
            action="cancellation_email_failed",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return False


def _send_reschedule_email(
    client: GraphClient,
    candidate_email: str,
    candidate_name: str,
    role_title: str,
    old_time: str,
    new_time: str,
    teams_url: Optional[str],
    company: CompanyConfig,
) -> bool:
    """Send reschedule notification email to candidate."""
    try:
        html_body = build_reschedule_email_html(
            candidate_name=candidate_name,
            role_title=role_title,
            old_time=old_time,
            new_time=new_time,
            teams_url=teams_url,
            company=company,
        )
        client.send_mail(
            subject=f"Interview Rescheduled: {role_title} at {company.name}",
            body=html_body,
            to_recipients=[candidate_email],
            content_type="HTML",
        )
        log_structured(
            LogLevel.INFO,
            f"Reschedule email sent to {candidate_email}",
            action="reschedule_email_sent",
        )
        return True
    except Exception as e:
        log_structured(
            LogLevel.ERROR,
            f"Failed to send reschedule email: {e}",
            action="reschedule_email_failed",
            error_type=type(e).__name__,
            exc_info=True,
        )
        return False


def _handle_reschedule(
    *,
    audit: AuditLog,
    event_id: str,
    new_date,
    new_time,
    duration_minutes: int,
    tz_name: str,
    context_row: Dict[str, Any],
    reason: str = "",
    notify_candidate: bool = True,
) -> None:
    """
    Handle interview reschedule with status tracking and notifications.

    Args:
        audit: AuditLog instance
        event_id: Graph event ID
        new_date: New interview date
        new_time: New interview time
        duration_minutes: Interview duration
        tz_name: Display timezone
        context_row: Interview data from database
        reason: Reason for reschedule
        notify_candidate: Whether to send email notification
    """
    client = _make_graph_client()
    if not client:
        st.error("Graph is not configured.")
        return

    # Calculate new times
    start_local = datetime.combine(new_date, new_time).replace(tzinfo=_zoneinfo(tz_name))
    end_local = start_local + timedelta(minutes=duration_minutes)
    start_utc = to_utc(start_local)
    end_utc = to_utc(end_local)

    # Store old time for notification - use candidate timezone for better UX
    old_time_str = context_row.get("start_utc", "")
    candidate_tz = context_row.get("candidate_timezone") or context_row.get("display_timezone") or tz_name
    old_time_formatted = _format_interview_time_for_candidate(old_time_str, candidate_tz, tz_name)

    # Format new time in candidate's timezone
    new_time_formatted = start_local.strftime("%A, %B %d, %Y at %I:%M %p") + f" ({tz_name})"

    patch = {
        "start": {"dateTime": start_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name},
        "end": {"dateTime": end_local.strftime("%Y-%m-%dT%H:%M:%S"), "timeZone": tz_name},
    }

    try:
        # Patch the Graph event
        client.patch_event(event_id, patch, send_updates="all")

        # Update interview status in database
        audit.update_interview_status(
            event_id=event_id,
            new_status=InterviewStatus.RESCHEDULED,
            reason=reason,
            updated_by=context_row.get("recruiter_email"),
        )

        # Increment ICS sequence for proper calendar client update
        audit.increment_ics_sequence(event_id)

        # Log success
        audit.log(
            "interview_rescheduled",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload={
                "old_start": old_time_str,
                "new_start": start_utc.isoformat(),
                "new_end": end_utc.isoformat(),
                "reason": reason,
                "notification_sent": notify_candidate,
            },
            status="success",
        )

        # Send notification email
        notification_sent = False
        if notify_candidate:
            company = get_company_config()
            candidate_email = context_row.get("candidate_email", "")
            if candidate_email:
                # Extract candidate name from context (candidates_json or fallback)
                candidate_name = _extract_candidate_name_from_context(context_row)
                notification_sent = _send_reschedule_email(
                    client=client,
                    candidate_email=candidate_email,
                    candidate_name=candidate_name,
                    role_title=context_row.get("role_title", ""),
                    old_time=old_time_formatted,
                    new_time=new_time_formatted,
                    teams_url=context_row.get("teams_join_url"),
                    company=company,
                )
                if notification_sent:
                    audit.log(
                        "notification_sent",
                        actor=context_row.get("recruiter_email", "") or "",
                        candidate_email=candidate_email,
                        event_id=event_id,
                        payload={"type": "reschedule"},
                        status="success",
                    )

        if notify_candidate and notification_sent:
            st.success("Event rescheduled. Candidate notified via email.")
        else:
            st.success("Event rescheduled. Attendees should receive updated invites.")

    except GraphAPIError as e:
        st.error("Reschedule failed.")
        st.json(e.response_json)
        audit.log(
            "interview_reschedule_failed",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload=e.response_json,
            status="failed",
            error_message=str(e),
        )


def _handle_cancel(
    *,
    audit: AuditLog,
    event_id: str,
    context_row: Dict[str, Any],
    reason: str = "",
    notify_candidate: bool = True,
    candidate_message: str = "",
) -> None:
    """
    Handle interview cancellation with status tracking and notifications.

    Args:
        audit: AuditLog instance
        event_id: Graph event ID
        context_row: Interview data from database
        reason: Cancellation reason
        notify_candidate: Whether to send email notification
        candidate_message: Optional custom message for candidate
    """
    client = _make_graph_client()
    if not client:
        st.error("Graph is not configured.")
        return

    # Format interview time for notification - use candidate timezone for better UX
    interview_time_str = context_row.get("start_utc", "")
    candidate_tz = context_row.get("candidate_timezone") or context_row.get("display_timezone")
    interview_time_formatted = _format_interview_time_for_candidate(interview_time_str, candidate_tz)

    try:
        # Delete the calendar event
        client.delete_event(event_id)

        # Update interview status in database
        audit.update_interview_status(
            event_id=event_id,
            new_status=InterviewStatus.CANCELLED,
            reason=reason,
            updated_by=context_row.get("recruiter_email"),
        )

        # Log success
        audit.log(
            "interview_cancelled",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload={
                "reason": reason,
                "notification_sent": notify_candidate,
            },
            status="success",
        )

        # Send notification email
        notification_sent = False
        if notify_candidate:
            company = get_company_config()
            candidate_email = context_row.get("candidate_email", "")
            if candidate_email:
                # Extract candidate name from context (candidates_json or fallback)
                candidate_name = _extract_candidate_name_from_context(context_row)
                notification_sent = _send_cancellation_email(
                    client=client,
                    candidate_email=candidate_email,
                    candidate_name=candidate_name,
                    role_title=context_row.get("role_title", ""),
                    interview_time=interview_time_formatted,
                    reason=reason,
                    custom_message=candidate_message,
                    company=company,
                )
                if notification_sent:
                    audit.log(
                        "notification_sent",
                        actor=context_row.get("recruiter_email", "") or "",
                        candidate_email=candidate_email,
                        event_id=event_id,
                        payload={"type": "cancellation", "reason": reason},
                        status="success",
                    )

        if notify_candidate and notification_sent:
            st.success("Interview cancelled. Candidate notified via email.")
        else:
            st.success("Interview cancelled. Attendees should receive cancellation notices.")

    except GraphAPIError as e:
        st.error("Cancel failed.")
        st.json(e.response_json)
        audit.log(
            "interview_cancel_failed",
            actor=context_row.get("recruiter_email", "") or "",
            candidate_email=context_row.get("candidate_email", "") or "",
            hiring_manager_email=context_row.get("hiring_manager_email", "") or "",
            recruiter_email=context_row.get("recruiter_email", "") or "",
            role_title=context_row.get("role_title", "") or "",
            event_id=event_id,
            payload=e.response_json,
            status="failed",
            error_message=str(e),
        )


if __name__ == "__main__":
    main()
